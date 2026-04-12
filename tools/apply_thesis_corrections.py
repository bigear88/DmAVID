#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fix_thesis_v1.py
依 CANONICAL_TRUTH.md 將論文章節 docx 內錯誤數字一次修正到位。

修正策略：
- INLINE：跨章節的字串替換（適用於段落與表格儲存格）。
- CELL_FIXES：限定章節 + 表格索引 + (row, col) 的精準替換，避免誤傷其他位置的相同字串。

每次執行前自動建立 .bak 備份。
"""

import shutil
from pathlib import Path
from docx import Document

REPO = Path(r"C:\Users\User\LLM-DEFI\DmAVID")

# ----------------------------------------------------------------------------
# 替換規則
# ----------------------------------------------------------------------------

# 跨章節的純字串替換：(old, new, comment)
INLINE = {
    "DmAVID_論文第三章.docx": [
        ("F1=0.8917", "F1=0.9061", "Ch3 baseline F1"),
        ("Precision=74.2%", "Precision=84.34%", "Ch3 baseline Precision"),
        ("Recall=98.6%", "Recall=97.90%", "Ch3 baseline Recall"),
        ("FPR=49%", "FPR=26%", "Ch3 baseline FPR"),
    ],
    "DmAVID_論文第四章.docx": [
        ("210/39", "25/39", "Ch4 typo: smart preprocess detected"),
        ("（+212%）", "（+108.3%）", "Ch4 wrong relative improvement"),
        ("+212%", "+108.3%", "Ch4 wrong relative (no parens)"),
        ("+43.59%", "+33.33pp", "Ch4 Table 4-20 smart preprocess improvement"),
        ("25.64%（10/39）", "7.69%（3/39）", "Ch4 EVMbench baseline (full-width)"),
        ("25.64% (10/39)", "7.69% (3/39)", "Ch4 EVMbench baseline (ascii)"),
        ("基準線的 25.64%", "基準線的 7.69%", "Ch4 baseline rate text"),
        ("12.82 個百分點", "23.08 個百分點", "Ch4 Enhanced absolute pp"),
        ("+166.67%", "+300.0%", "Ch4 Enhanced relative improvement"),
        ("新增偵測到 5 個", "新增偵測到 9 個", "Ch4 new detections count"),
        # 第二輪補充修正
        ("LLM+RAG 基準線 / 10/39 (25.64%) / +25.64%",
         "LLM+RAG 基準線 / 3/39 (7.69%) / +7.69pp",
         "Ch4 Table 4-20 (text) LLM+RAG baseline row"),
        ("EVMbench 上的 25.64% 偵測率",
         "EVMbench 上的 7.69% 偵測率",
         "Ch4 p146 EVMbench rate"),
        ("成功偵測到 3 個漏洞（偵測率 25.64%）",
         "成功偵測到 3 個漏洞（偵測率 7.69%）",
         "Ch4 p156/p201 first occurrence"),
        ("同樣偵測到 3 個漏洞（25.64%）",
         "同樣偵測到 3 個漏洞（7.69%）",
         "Ch4 p201 Hybrid mention"),
        ("本研究的 LLM 方法（25.64%）",
         "本研究的 LLM 方法（7.69%）",
         "Ch4 p211 LLM method rate"),
        ("均為 10/39", "均為 3/39", "Ch4 p210 Hybrid vs LLM+RAG"),
        ("僅成功偵測到 3 個（偵測率 25.64%）",
         "僅成功偵測到 3 個（偵測率 7.69%）",
         "Ch4 p156 alt phrasing"),
        # Table 4-5 敘述段 (p70): 原文寫 LLM+RAG 在 1x-20x 最低 — 連舊表的「最佳方法」欄都標 Hybrid，自相矛盾
        ("分析結果顯示，在絕大多數成本比率（1x 至 20x）下，LLM+RAG 都是總成本最低的方法。",
         "分析結果顯示，在絕大多數成本比率（1x 至 20x）下，混合式框架（LLM+RAG+Self-Verify）都是總成本最低的方法。",
         "Ch4 p70 Table 4-5 narrative"),
        # LLM Base FPR 應為 95%（FP=95, TN=5），原文寫 98%
        ("但代價是 98% 的誤報率",
         "但代價是 95% 的誤報率（FPR=0.95）",
         "Ch4 p70 LLM Base FPR"),
    ],
    "DmAVID_論文第五章.docx": [
        ("F1 從 0.8917 提升至 0.8924（+5.4%）",
         "F1 從 0.8917 提升至 0.8924（+0.08%）",
         "Ch5 iteration improvement (CRITICAL)"),
        ("F1 相對提升 +20.6%", "F1 相對提升 +21.2%", "Ch5 ablation RAG contribution"),
        # 為 0.8917 加註，說明此為 self-verify 子實驗 baseline，與主管線 0.9061 不同
        ("經 3 輪對抗迭代後 F1 從 0.8917 提升至 0.8924（+0.08%）",
         "經 3 輪對抗迭代後 F1 從 0.8917 提升至 0.8924（+0.08%；註：此為 Self-Verify 子實驗 baseline，主管線 LLM+RAG+Self-Verify 之 F1=0.9121）",
         "Ch5 p22 add baseline footnote"),
        ("多代理迭代進一步提升（+0.08%，從 0.8917→0.8924）",
         "多代理迭代進一步提升（+0.08%，self-verify 子實驗：從 0.8917→0.8924）",
         "Ch5 p23 ablation contribution clarification"),
    ],
    # 主合併 docx：只有少量 Ch3 內容
    "DmAVID基於多代理對抗式迭代之以太坊 DeFi 智能合約漏洞偵測研究.docx": [
        ("F1=0.8917", "F1=0.9061", "main combined: Ch3 baseline F1"),
        ("Precision=74.2%", "Precision=84.34%", "main combined: Ch3 baseline Precision"),
        ("Recall=98.6%", "Recall=97.90%", "main combined: Ch3 baseline Recall"),
        ("FPR=49%", "FPR=26%", "main combined: Ch3 baseline FPR"),
    ],
}

# 表格儲存格座標精準修正：(chapter, table_idx, row, col, old, new, comment)
CELL_FIXES = [
    # Table 4-5 (idx 4): 成本敏感分析 — 用 canonical FP/FN 重算
    # 公式: cost = FP + FN × ratio
    # 舊表用了 LLM+RAG FP=57/FN=1 與 Hybrid FP=52/FN=1（不存在於 GitHub 的更早版本）
    # 新值: Slither (FP=84,FN=8) 不變; LLM Base (FP=95,FN=1); LLM+RAG (FP=26,FN=3); Hybrid (FP=24,FN=3)
    ("DmAVID_論文第四章.docx", 4, 1, 2, "95", "96", "Table 4-5 LLM Base 1x"),
    ("DmAVID_論文第四章.docx", 4, 2, 2, "95", "97", "Table 4-5 LLM Base 2x"),
    ("DmAVID_論文第四章.docx", 4, 3, 2, "95", "100", "Table 4-5 LLM Base 5x"),
    ("DmAVID_論文第四章.docx", 4, 4, 2, "95", "105", "Table 4-5 LLM Base 10x"),
    ("DmAVID_論文第四章.docx", 4, 5, 2, "95", "115", "Table 4-5 LLM Base 20x"),
    ("DmAVID_論文第四章.docx", 4, 6, 2, "95", "145", "Table 4-5 LLM Base 50x"),
    ("DmAVID_論文第四章.docx", 4, 1, 3, "58", "29", "Table 4-5 LLM+RAG 1x"),
    ("DmAVID_論文第四章.docx", 4, 2, 3, "59", "32", "Table 4-5 LLM+RAG 2x"),
    ("DmAVID_論文第四章.docx", 4, 3, 3, "62", "41", "Table 4-5 LLM+RAG 5x"),
    ("DmAVID_論文第四章.docx", 4, 4, 3, "67", "56", "Table 4-5 LLM+RAG 10x"),
    ("DmAVID_論文第四章.docx", 4, 5, 3, "77", "86", "Table 4-5 LLM+RAG 20x"),
    ("DmAVID_論文第四章.docx", 4, 6, 3, "107", "176", "Table 4-5 LLM+RAG 50x"),
    ("DmAVID_論文第四章.docx", 4, 1, 4, "53", "27", "Table 4-5 Hybrid 1x"),
    ("DmAVID_論文第四章.docx", 4, 2, 4, "54", "30", "Table 4-5 Hybrid 2x"),
    ("DmAVID_論文第四章.docx", 4, 3, 4, "57", "39", "Table 4-5 Hybrid 5x"),
    ("DmAVID_論文第四章.docx", 4, 4, 4, "62", "54", "Table 4-5 Hybrid 10x"),
    ("DmAVID_論文第四章.docx", 4, 5, 4, "72", "84", "Table 4-5 Hybrid 20x"),
    ("DmAVID_論文第四章.docx", 4, 6, 4, "102", "174", "Table 4-5 Hybrid 50x"),

    # Table 4-1 (idx 0): 與相關研究比較 — 混合式框架 row
    ("DmAVID_論文第四章.docx", 0, 4, 2, "0.73", "0.85", "Table 4-1 混合式框架 Precision"),
    ("DmAVID_論文第四章.docx", 0, 4, 3, "0.99", "0.98", "Table 4-1 混合式框架 Recall"),
    ("DmAVID_論文第四章.docx", 0, 4, 4, "0.84", "0.91", "Table 4-1 混合式框架 F1"),

    # Table 4-3 (idx 2): 消融比較 — 僅 LLM Recall
    ("DmAVID_論文第四章.docx", 2, 2, 2, "1.00", "0.99", "Table 4-3 僅 LLM Recall"),

    # Table 4-11 (idx 10): 階段消融 — RAG/Self-Verify 相對提升
    ("DmAVID_論文第四章.docx", 10, 3, 2, "+20.6%", "+21.2%", "Table 4-11 RAG contribution"),
    ("DmAVID_論文第四章.docx", 10, 4, 2, "+6.6%", "+0.66%", "Table 4-11 Self-Verify contribution"),

    # Table 4-13 (idx 12): Agent 貢獻 — RAG/Self-Verify 相對提升
    ("DmAVID_論文第四章.docx", 12, 2, 1, "+20.6%", "+21.2%", "Table 4-13 RAG contribution"),
    ("DmAVID_論文第四章.docx", 12, 3, 1, "+6.6%", "+0.66%", "Table 4-13 Self-Verify contribution"),

    # Table 4-14 (idx 13): EVMbench 專案類別 — 總計列偵測率
    ("DmAVID_論文第四章.docx", 13, 5, 3, "25.64%", "7.69%", "Table 4-14 總計 偵測率"),

    # Table 4-15 (idx 14): EVMbench 方法比較 — LLM+RAG / Hybrid 行
    ("DmAVID_論文第四章.docx", 14, 3, 1, "10/39", "3/39", "Table 4-15 LLM+RAG detected"),
    ("DmAVID_論文第四章.docx", 14, 3, 2, "25.64%", "7.69%", "Table 4-15 LLM+RAG rate"),
    ("DmAVID_論文第四章.docx", 14, 4, 1, "10/39", "3/39", "Table 4-15 Hybrid detected"),
    ("DmAVID_論文第四章.docx", 14, 4, 2, "25.64%", "7.69%", "Table 4-15 Hybrid rate"),

    # Table 4-17 (idx 16): Enhanced 比較
    ("DmAVID_論文第四章.docx", 16, 1, 2, "8 / 39", "12 / 39", "Table 4-17 Enhanced detected"),
    ("DmAVID_論文第四章.docx", 16, 1, 3, "+5", "+9", "Table 4-17 detected delta"),
    ("DmAVID_論文第四章.docx", 16, 2, 1, "25.64%", "7.69%", "Table 4-17 baseline rate"),
    ("DmAVID_論文第四章.docx", 16, 2, 3, "+166.67%", "+300.0%", "Table 4-17 relative improvement"),
    ("DmAVID_論文第四章.docx", 16, 3, 3, "+12.82 pp", "+23.08 pp", "Table 4-17 absolute improvement"),
    ("DmAVID_論文第四章.docx", 16, 4, 3, "+5 個", "+9 個", "Table 4-17 new detections (text)"),
]

# ----------------------------------------------------------------------------
# 工具函式
# ----------------------------------------------------------------------------

def replace_in_paragraph(para, old, new):
    """段落字串替換，盡量保留原 run 格式。"""
    if old not in para.text:
        return False
    # 嘗試 single-run 替換（保留格式）
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # 跨 run：合併到第一個 run
    new_full = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_in_cell(cell, old, new):
    for para in cell.paragraphs:
        if replace_in_paragraph(para, old, new):
            return True
    return False


def apply_inline_to_doc(doc, old, new):
    """對整份 doc 的所有段落與表格儲存格做字串替換，回傳命中次數。"""
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            replace_in_paragraph(para, old, new)
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    replace_in_cell(cell, old, new)
                    count += 1
    return count


# ----------------------------------------------------------------------------
# 主流程
# ----------------------------------------------------------------------------

def main():
    log = []
    chapters = sorted(set(list(INLINE.keys()) + [c[0] for c in CELL_FIXES]))

    for chapter in chapters:
        path = REPO / chapter
        if not path.exists():
            log.append(f"SKIP missing: {chapter}")
            continue

        # 備份
        bak = path.with_suffix(path.suffix + ".bak")
        if not bak.exists():
            shutil.copy2(path, bak)
            log.append(f"  [{chapter}] backup -> {bak.name}")

        doc = Document(str(path))
        modified = 0

        # 1. INLINE 替換
        for old, new, comment in INLINE.get(chapter, []):
            count = apply_inline_to_doc(doc, old, new)
            if count > 0:
                log.append(f"  [{chapter}] INLINE x{count} | {comment}: '{old}' -> '{new}'")
                modified += count
            else:
                log.append(f"  [{chapter}] WARN INLINE 0 hit | {comment}: '{old}'")

        # 2. CELL 座標精準修正
        for cf_chap, ti, ri, ci, old, new, comment in CELL_FIXES:
            if cf_chap != chapter:
                continue
            if ti >= len(doc.tables):
                log.append(f"  [{chapter}] SKIP tbl{ti} out of range")
                continue
            table = doc.tables[ti]
            if ri >= len(table.rows):
                log.append(f"  [{chapter}] SKIP tbl{ti}r{ri} out of range")
                continue
            cells = table.rows[ri].cells
            if ci >= len(cells):
                log.append(f"  [{chapter}] SKIP tbl{ti}r{ri}c{ci} out of range")
                continue
            cell = cells[ci]
            before = cell.text
            if old not in before:
                log.append(f"  [{chapter}] WARN CELL tbl{ti}r{ri}c{ci} | {comment}: '{old}' not in '{before[:40]}'")
                continue
            replace_in_cell(cell, old, new)
            log.append(f"  [{chapter}] CELL tbl{ti}r{ri}c{ci} | {comment}: '{old}' -> '{new}'")
            modified += 1

        doc.save(str(path))
        print(f"\n[OK] {chapter}: {modified} modifications saved")

    # 輸出 log
    log_path = REPO / "THESIS_FIX_LOG.txt"
    log_path.write_text("\n".join(log), encoding="utf-8")
    print(f"\nLog written: {log_path}")
    print(f"Total log lines: {len(log)}")


if __name__ == "__main__":
    main()
