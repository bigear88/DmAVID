#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
validate_thesis_tables.py

Auto-validate all numeric claims in DmAVID thesis chapters against canonical
experimental results. Output: THESIS_DIFF_REPORT.md

Usage:
    python tools/validate_thesis_tables.py

Canonical truth lives in CANONICAL_TRUTH.md (human-readable) and the JSON files
loaded below (machine-readable).
"""

import json
import os
import re
import sys
from pathlib import Path
from docx import Document

REPO = Path(__file__).resolve().parent.parent
WSL_REPO = Path(r"\\wsl.localhost\Ubuntu\home\curtis\DmAVID")

CHAPTER_FILES = [
    REPO / "DmAVID_論文第一章.docx",
    REPO / "DmAVID_論文第二章.docx",
    REPO / "DmAVID_論文第三章.docx",
    REPO / "DmAVID_論文第四章.docx",
    REPO / "DmAVID_論文第五章.docx",
]

REPORT_PATH = REPO / "THESIS_DIFF_REPORT.md"


def load_canonical():
    """Load all canonical JSON results from WSL repo."""
    base = WSL_REPO / "experiments"
    files = {
        "slither":            base / "slither" / "slither_results.json",
        "llm_rag":            base / "llm_rag" / "llm_rag_results.json",
        "hybrid":             base / "hybrid" / "hybrid_results.json",
        "self_verify":        base / "hybrid" / "self_verify_results.json",
        "ablation_v5_clean":  base / "ablation" / "ablation_v5_clean_results.json",
        "evmbench_llm_rag":   base / "evmbench" / "evmbench_detect_results.json",
        "evmbench_hybrid":    base / "evmbench" / "evmbench_hybrid_results.json",
        "evmbench_enhanced":  base / "evmbench_enhanced" / "enhanced_results.json",
        "evmbench_smart":     base / "evmbench_smart" / "smart_preprocess_results.json",
        "defi_real_world":    base / "defi_real_world" / "defi_results.json",
    }
    out = {}
    for k, p in files.items():
        if p.exists():
            with open(p, "r", encoding="utf-8") as f:
                out[k] = json.load(f)
        else:
            print(f"WARN: missing {p}")
    return out


def build_truth_table(canon):
    """Flatten canonical data into a list of (label, value, source) facts to check."""
    facts = []

    def add(label, value, source, kind="number", tol=0.001):
        facts.append({
            "label": label,
            "value": value,
            "source": source,
            "kind": kind,
            "tol": tol,
        })

    # SmartBugs Slither
    m = canon["slither"]["metrics"]
    add("Slither TP",        m["tp"], "slither_results.json")
    add("Slither FP",        m["fp"], "slither_results.json")
    add("Slither FN",        m["fn"], "slither_results.json")
    add("Slither TN",        m["tn"], "slither_results.json")
    add("Slither Precision", m["precision"], "slither_results.json")
    add("Slither Recall",    m["recall"], "slither_results.json")
    add("Slither F1",        m["f1_score"], "slither_results.json")

    # SmartBugs LLM+RAG (canonical baseline)
    m = canon["llm_rag"]["metrics"]
    add("LLM+RAG TP",        m["tp"], "llm_rag_results.json")
    add("LLM+RAG FP",        m["fp"], "llm_rag_results.json")
    add("LLM+RAG FN",        m["fn"], "llm_rag_results.json")
    add("LLM+RAG TN",        m["tn"], "llm_rag_results.json")
    add("LLM+RAG Precision", m["precision"], "llm_rag_results.json")
    add("LLM+RAG Recall",    m["recall"], "llm_rag_results.json")
    add("LLM+RAG F1",        m["f1_score"], "llm_rag_results.json")

    # Hybrid (Slither+LLM+RAG fusion)
    m = canon["hybrid"]["metrics"]
    add("Hybrid TP",         m["tp"], "hybrid_results.json")
    add("Hybrid FP",         m["fp"], "hybrid_results.json")
    add("Hybrid FN",         m["fn"], "hybrid_results.json")
    add("Hybrid TN",         m["tn"], "hybrid_results.json")
    add("Hybrid Precision",  m["precision"], "hybrid_results.json")
    add("Hybrid Recall",     m["recall"], "hybrid_results.json")
    add("Hybrid F1",         m["f1_score"], "hybrid_results.json")

    # Ablation v5_clean (each config)
    for cfg in canon["ablation_v5_clean"]["configs"]:
        name = cfg["config"]
        cm = cfg["metrics"]
        add(f"Ablation[{name}] TP", cm["tp"], "ablation_v5_clean_results.json")
        add(f"Ablation[{name}] FP", cm["fp"], "ablation_v5_clean_results.json")
        add(f"Ablation[{name}] FN", cm["fn"], "ablation_v5_clean_results.json")
        add(f"Ablation[{name}] TN", cm["tn"], "ablation_v5_clean_results.json")
        add(f"Ablation[{name}] F1", cm["f1"], "ablation_v5_clean_results.json")

    # EVMbench
    add("EVMbench LLM+RAG detected",   canon["evmbench_llm_rag"]["total_detected"],
        "evmbench_detect_results.json")
    add("EVMbench Hybrid detected",    canon["evmbench_hybrid"]["total_detected"],
        "evmbench_hybrid_results.json")
    add("EVMbench Enhanced detected",  canon["evmbench_enhanced"]["overall"]["total_detected"],
        "evmbench_enhanced/enhanced_results.json")
    add("EVMbench Enhanced rate",      canon["evmbench_enhanced"]["overall"]["detect_rate"],
        "evmbench_enhanced/enhanced_results.json")
    add("EVMbench Smart detected",     canon["evmbench_smart"]["total_detected"],
        "evmbench_smart/smart_preprocess_results.json")
    add("EVMbench Smart rate",         canon["evmbench_smart"]["detect_rate"],
        "evmbench_smart/smart_preprocess_results.json")
    add("EVMbench total gold",         canon["evmbench_smart"]["total_gold"],
        "evmbench_smart/smart_preprocess_results.json")

    # DeFi real-world LLM+RAG
    m = canon["defi_real_world"]["llm_rag"]
    add("DeFi LLM+RAG TP",        m["tp"], "defi_real_world/defi_results.json")
    add("DeFi LLM+RAG FP",        m["fp"], "defi_real_world/defi_results.json")
    add("DeFi LLM+RAG FN",        m["fn"], "defi_real_world/defi_results.json")
    add("DeFi LLM+RAG TN",        m["tn"], "defi_real_world/defi_results.json")
    add("DeFi LLM+RAG F1",        m["f1"], "defi_real_world/defi_results.json")
    add("DeFi LLM+RAG Recall",    m["recall"], "defi_real_world/defi_results.json")

    return facts


def extract_thesis_text(docx_path):
    """Pull all paragraph + table cell text from a chapter docx."""
    doc = Document(str(docx_path))
    chunks = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(("para", t))
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text.strip()
                if t:
                    chunks.append((f"tbl{ti}r{ri}c{ci}", t))
    return chunks


# Numeric extraction patterns
NUM_INT  = re.compile(r"(?<![\d.])(\d{1,4})(?![\d.])")
NUM_PCT  = re.compile(r"(\d+(?:\.\d+)?)\s*%")
NUM_DEC  = re.compile(r"\b(0?\.\d{2,4})\b")
FRAC     = re.compile(r"\b(\d{1,3})\s*/\s*(\d{1,4})\b")


def extract_numbers(text):
    """Return set of (kind, value) pairs found in text."""
    out = set()
    for m in NUM_PCT.finditer(text):
        out.add(("pct", round(float(m.group(1)) / 100.0, 4)))
    for m in NUM_DEC.finditer(text):
        out.add(("dec", round(float(m.group(1)), 4)))
    for m in NUM_INT.finditer(text):
        v = int(m.group(1))
        if 0 < v < 10000:
            out.add(("int", v))
    for m in FRAC.finditer(text):
        num = int(m.group(1))
        den = int(m.group(2))
        if den > 0:
            out.add(("frac", (num, den)))
    return out


def find_fact_in_chunks(fact, chunks):
    """Locate where this fact's value appears (or doesn't) in chapter chunks."""
    val = fact["value"]
    tol = fact["tol"]
    hits = []
    for loc, text in chunks:
        nums = extract_numbers(text)
        if isinstance(val, float):
            if val < 1.0:
                target_dec = round(val, 4)
                target_pct = round(val, 4)
                for kind, v in nums:
                    if kind in ("dec", "pct") and abs(v - target_dec) <= tol:
                        hits.append(loc)
                        break
            else:
                for kind, v in nums:
                    if kind in ("dec", "int") and abs(v - val) <= tol:
                        hits.append(loc)
                        break
        elif isinstance(val, int):
            for kind, v in nums:
                if kind == "int" and v == val:
                    hits.append(loc)
                    break
    return hits


# Wrong/suspect numbers we already KNOW are bad — flag with explanation
KNOWN_BAD = [
    {
        "pattern": re.compile(r"210\s*/\s*39"),
        "issue": "Typo: '210/39' should be '25/39' (Smart preprocess result)",
        "fix": "25/39 = 64.10%",
    },
    {
        "pattern": re.compile(r"\+\s*212\s*%"),
        "issue": "Wrong relative improvement",
        "fix": "(64.10-30.77)/30.77 = +108.3%, or (64.10-7.69)/7.69 = +733%",
    },
    {
        "pattern": re.compile(r"\+\s*5\.4\s*%"),
        "issue": "Wrong improvement claim for adversarial iteration",
        "fix": "F1 0.8917→0.8924 = +0.08% relative, or +0.07pp absolute",
    },
    {
        "pattern": re.compile(r"0\.8917"),
        "issue": "Older LLM+RAG snapshot, NOT canonical baseline",
        "fix": "Replace with 0.9061 (canonical baseline from llm_rag_results.json)",
    },
    {
        "pattern": re.compile(r"\+\s*43\.59\s*%"),
        "issue": "Wrong improvement column in Table 4-20",
        "fix": "+33.33pp (absolute) or +108.3% (relative)",
    },
    {
        "pattern": re.compile(r"\+\s*25\.64\s*%"),
        "issue": "Wrong improvement value in Table 4-20 iterative row (copy-paste from above)",
        "fix": "+5.13pp (absolute)",
    },
    {
        "pattern": re.compile(r"10\s*/\s*39"),
        "issue": "EVMbench LLM+RAG should be 3/39, not 10/39",
        "fix": "3/39 = 7.69%",
    },
    {
        "pattern": re.compile(r"8\s*/\s*39"),
        "issue": "EVMbench Enhanced should be 12/39, not 8/39",
        "fix": "12/39 = 30.77%",
    },
    {
        "pattern": re.compile(r"gpt[- ]?5\.4[- ]?mini", re.IGNORECASE),
        "issue": "Non-existent model name",
        "fix": "gpt-4.1-mini",
    },
]


# 白名單：含這些上下文的 0.8917 是 user-approved（Self-Verify 子實驗 baseline，Option A）
ALLOWED_0_8917_CONTEXTS = [
    "Self-Verify 子實驗",
    "self-verify 子實驗",
]


def scan_known_bad(chapter_name, chunks):
    """Find every line containing a known-bad pattern."""
    findings = []
    for loc, text in chunks:
        for kb in KNOWN_BAD:
            if kb["pattern"].search(text):
                # 0.8917 白名單例外：在含有 self-verify 子實驗註解的段落中允許
                if r"0\.8917" in kb["pattern"].pattern and any(c in text for c in ALLOWED_0_8917_CONTEXTS):
                    continue
                findings.append({
                    "chapter": chapter_name,
                    "location": loc,
                    "snippet": text[:200],
                    "issue": kb["issue"],
                    "fix": kb["fix"],
                })
    return findings


def main():
    print("=" * 70)
    print("DmAVID Thesis Validator")
    print("=" * 70)

    print("\n[1/4] Loading canonical truth from JSON...")
    canon = load_canonical()
    print(f"  loaded: {list(canon.keys())}")

    print("\n[2/4] Building truth table...")
    facts = build_truth_table(canon)
    print(f"  {len(facts)} facts to verify")

    print("\n[3/4] Reading chapter docx files...")
    all_chunks = {}
    for cf in CHAPTER_FILES:
        if not cf.exists():
            print(f"  MISSING: {cf.name}")
            continue
        chunks = extract_thesis_text(cf)
        all_chunks[cf.name] = chunks
        print(f"  {cf.name}: {len(chunks)} text blocks")

    print("\n[4/4] Validating...\n")

    # Combined chunks across all chapters
    combined = []
    for fname, chunks in all_chunks.items():
        for loc, text in chunks:
            combined.append((f"{fname}::{loc}", text))

    # 1) Check whether each canonical fact appears anywhere
    missing_facts = []
    found_facts = []
    for fact in facts:
        hits = find_fact_in_chunks(fact, combined)
        if hits:
            found_facts.append((fact, hits[:3]))
        else:
            missing_facts.append(fact)

    # 2) Scan for known-bad patterns
    bad_findings = []
    for fname, chunks in all_chunks.items():
        bad_findings.extend(scan_known_bad(fname, chunks))

    # Write report
    lines = []
    lines.append("# DmAVID Thesis Diff Report")
    lines.append("")
    lines.append(f"Generated by `tools/validate_thesis_tables.py`")
    lines.append("")
    lines.append("Canonical source: see `CANONICAL_TRUTH.md` and `experiments/**/*.json`")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append(f"## Summary")
    lines.append("")
    lines.append(f"- Canonical facts checked: **{len(facts)}**")
    lines.append(f"- Facts found in thesis text: **{len(found_facts)}**")
    lines.append(f"- Facts NOT found in thesis text: **{len(missing_facts)}**")
    lines.append(f"- Known-bad patterns detected: **{len(bad_findings)}**")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append("## A. Known-bad patterns (CRITICAL — must fix)")
    lines.append("")
    if not bad_findings:
        lines.append("_None detected._")
    else:
        for i, b in enumerate(bad_findings, 1):
            lines.append(f"### {i}. [{b['chapter']}] {b['issue']}")
            lines.append(f"- **Location:** `{b['location']}`")
            lines.append(f"- **Snippet:** `{b['snippet']}`")
            lines.append(f"- **Fix:** {b['fix']}")
            lines.append("")
    lines.append("---")
    lines.append("")

    lines.append("## B. Canonical facts NOT found in thesis (may indicate omission OR wrong value)")
    lines.append("")
    if not missing_facts:
        lines.append("_All canonical facts located in thesis text._")
    else:
        lines.append("| # | Label | Canonical Value | Source File |")
        lines.append("|---|---|---|---|")
        for i, f in enumerate(missing_facts, 1):
            v = f["value"]
            if isinstance(v, float):
                vs = f"{v:.4f}" if v < 1 else f"{v:.4f}"
            else:
                vs = str(v)
            lines.append(f"| {i} | {f['label']} | {vs} | {f['source']} |")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append("## C. Canonical facts found in thesis (informational)")
    lines.append("")
    lines.append("| Label | Value | First location |")
    lines.append("|---|---|---|")
    for f, hits in found_facts:
        v = f["value"]
        vs = f"{v:.4f}" if isinstance(v, float) else str(v)
        loc = hits[0] if hits else "?"
        lines.append(f"| {f['label']} | {vs} | `{loc[:60]}` |")
    lines.append("")

    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")
    print(f"\nReport written: {REPORT_PATH}")
    print(f"  found:    {len(found_facts)}")
    print(f"  missing:  {len(missing_facts)}")
    print(f"  bad:      {len(bad_findings)}")


if __name__ == "__main__":
    main()
