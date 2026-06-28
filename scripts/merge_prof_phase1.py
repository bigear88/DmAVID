#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Phase 1: merge professor's accepted tracked-changes (Ch1-3 + Ch4-intro) into the
working doc by exact full-paragraph match+replace. Skips reference entries (orig empty,
handled in Phase 3). Professor edits are in paragraphs not touched by my Ch4 edits."""
import json, shutil
from docx import Document

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_prof_merge.docx"
shutil.copyfile(SRC, BAK); print("backup:", BAK)
pairs = json.load(open("/tmp/prof_pairs.json", encoding="utf-8"))
doc = Document(SRC)


def set_text(p, txt):
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)


# index my paragraphs by stripped text
by_text = {}
for p in doc.paragraphs:
    by_text.setdefault(p.text.strip(), []).append(p)

applied = miss = skipped = 0
for i, pr in enumerate(pairs):
    o = pr["orig"].strip(); a = pr["acc"].strip()
    if not o:                       # reference inserts -> Phase 3
        skipped += 1; continue
    if o == a:
        continue
    hits = by_text.get(o)
    if hits:
        set_text(hits[0], a)
        # update index so we don't re-hit
        by_text[o].pop(0)
        applied += 1
        print(f"[OK {i}] {a[:46]}...")
    else:
        miss += 1
        print(f"[MISS {i}] orig not found: {o[:50]}")

doc.save(SRC)
print(f"\nDONE applied={applied} miss={miss} skipped(ref)={skipped}")
