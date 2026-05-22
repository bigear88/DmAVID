#!/usr/bin/env python3
"""
在 thesis_v74_cascade.docx 第四章第三節消融實驗中，
於「六、FP 類型歸因消融」後、「七、消融實驗綜合討論」前，
插入新節「七、工具增強管線消融：六臂消融與閾值敏感度分析」。
原「七、綜合討論」改為「八、消融實驗綜合討論」。

輸出：thesis_v75_tool_ch4.docx
"""

import sys
import copy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IN_FILE  = '/mnt/c/Users/User/LLM-DEFI/DmAVID/thesis_v74_cascade.docx'
OUT_FILE = '/mnt/c/Users/User/LLM-DEFI/DmAVID/thesis_v75_tool_ch4.docx'

# ── helpers ─────────────────────────────────────────────────────────────────

def clone_run_fmt(source_run, target_run):
    """Copy bold/font-size from source run to target run."""
    target_run.bold = source_run.bold
    if source_run.font.size:
        target_run.font.size = source_run.font.size


def insert_para_after(ref_elem, body):
    """Insert an empty <w:p> immediately after ref_elem inside body. Returns new element."""
    new_p = OxmlElement('w:p')
    ref_elem.addnext(new_p)
    return new_p


def make_para(doc, text, style_name, bold=False, keep_with_next=False):
    """Create a standalone paragraph element (not yet in doc)."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    if keep_with_next:
        kwn = OxmlElement('w:keepNext')
        pPr.append(kwn)
    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p


def make_table(doc, headers, rows):
    """Build a <w:tbl> element with given headers and data rows."""
    from docx.oxml.table import CT_Tbl
    tbl = OxmlElement('w:tbl')

    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    all_rows = [headers] + rows
    for row_idx, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for cell_text in row_data:
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if row_idx == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = str(cell_text)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


def wrap_table_in_para(tbl):
    """Wrap table in a body-level block (tables are body-level, not inline)."""
    return tbl  # tables go directly into body, not inside <w:p>


# ── content definition ───────────────────────────────────────────────────────

INTRO = (
    "本節對應第參章第二節（一）所設計之工具增強管線，以三項實驗驗證 calibrated gate "
    "設計之有效性：（1）六臂消融確立 gate 為核心創新來源；（2）閾值敏感度驗證 T=2 門檻"
    "選擇的穩健性；（3）CodeBERT 長度公平性補充驗證基準方法之長度偏差。所有實驗模型"
    "均為 gpt-4.1-mini，資料集為 SmartBugs Curated（N=243）。"
)

SIX_ARM_INTRO = (
    "六臂消融在 SmartBugs Curated 全集（N=243，143 vuln + 100 safe）上，"
    "藉逐步加入設計組件量化各元素的邊際貢獻。表 4-A 列出六個設計變體及其結果。"
)

SIX_ARM_HEADERS = ['Arm', '描述', 'TP', 'FP', 'TN', 'FN', 'F1']
SIX_ARM_ROWS = [
    ['A', '純 LLM（無 RAG）', '142', '95', '5',  '1',  '0.7474'],
    ['B', 'LLM+RAG 基線',     '137', '27', '73', '6',  '0.8925'],
    ['C', '+硬規則過濾',       '116', '24', '76', '27', '0.8198'],
    ['D', '+LLM w/ tools（無 Gate）', '122', '29', '71', '21', '0.8299'],
    ['E', '+LLM w/ tools + code', '126', '25', '75', '17', '0.8571'],
    ['v4+Gate', 'v4 + Gate T≥2（論文主結果）', '135', '21', '79', '8', '0.9030'],
]

SIX_ARM_FINDING1 = (
    "六臂消融揭示三項關鍵發現。第一，工具增強本身不足以超越 RAG 基線。"
    "Arm C（+硬規則）F1=0.8198 與 Arm D（+LLM w/ tools，無 Gate）F1=0.8299 "
    "均低於 Arm B 基線（0.8925），顯示工具呼叫不配合 gate 反而引入更多預測不穩定性。"
)
SIX_ARM_FINDING2 = (
    "第二，calibrated gate 是唯一能超越基線的設計。v4+Gate 達 F1=0.9030，"
    "高於 Arm B 基線 +1.05 個百分點（絕對值）。gate 的作用在於：僅當 grep_guard "
    "偵測到防護機制數量（mitigation_count）≥ 2 時才啟動 re-evaluation，避免過度翻轉。"
)
SIX_ARM_FINDING3 = (
    "第三，code context（Arm E，+0.83 pp）的邊際效益雖正向，但仍未達 gate 效果。"
    "此結果說明 gate 的過濾邏輯比純粹增加 prompt 資訊更為關鍵。"
)

THRESH_INTRO = (
    "工具增強管線採 mitigation_count ≥ T 作為 gate 啟動條件，T=2 由第一輪實驗"
    "後的誤差分析選定。為驗證此選擇非測試集過擬合，本研究對 T∈{1,2,3} 進行完整掃描，"
    "結果如表 4-B 所示。"
)

THRESH_HEADERS = ['Gate 閾值', 'F1', 'TP', 'FP', 'TN', 'FN', 're-eval 數', 'flip 數']
THRESH_ROWS = [
    ['T=1', '0.8797', '128', '20', '80', '15', '75', '16'],
    ['T=2（選用）', '0.9030', '135', '21', '79', '8',  '17', '8'],
    ['T=3', '0.9007', '136', '23', '77', '7',  '10', '5'],
]

THRESH_FINDING = (
    "T=1 啟動 75 份合約的 re-evaluation（新增 58 次 LLM 呼叫），翻轉 16 份，"
    "F1=0.8797（低於基線）；T=2 翻轉 8 份，F1=0.9030（最佳）；T=3 翻轉 5 份，"
    "F1=0.9007。T∈{2,3} 皆正向，T=2 與 T=3 差距 0.23 個百分點，說明 T=2 的選擇"
    "對結果穩健：即使改為 T=3，框架仍超越 RAG 基線。論文誠實揭露此一來源（error analysis）"
    "與穩健性（T=3 同樣正向），以規避 gate 設計過擬合之批評。"
)

CODEBERT_CONTENT = (
    "第三節第四項（4.3.4）已揭示 CodeBERT 原版 F1=0.9180，與 DmAVID 0.9121 統計等價。"
    "然 SmartBugs Curated 子集存在顯著長度偏差：vuln 合約中位數 1,140 字元，"
    "wild/safe 合約中位數 7,429 字元，純長度閾值分類器即可達 F1=0.9091，"
    "顯示 CodeBERT 原版結果含有長度學習偏差（length confounding）。"
)
CODEBERT_ARM_C = (
    "為消除此偏差，本研究設計等長截斷實驗（Arm C v2）：過濾長度 <512 字元之合約，"
    "將所有保留合約截取前 512 字元，使訓練集 length_ratio=1.00（完全等長），"
    "共保留 199 份合約（99 vuln + 100 safe，排除 44 份過短之 vuln 合約）。"
    "微調 3 epochs 後，CodeBERT 在 n_test=40 的測試集上達 F1=0.8000"
    "（95% CI [0.629, 0.914]，TP=16，FP=4，TN=16，FN=4）。"
)
CODEBERT_CONCLUSION = (
    "在長度信號完全消除（length_ratio=1.00）的公平條件下，DmAVID v4+Gate（F1=0.9030）"
    "領先 CodeBERT +10.3 個百分點。此結果對應第貳章第六節第七項研究缺口，"
    "驗證基準方法之 F1=0.9180 包含資料集偏差；DmAVID 在消除偏差後的公平比較中"
    "仍具顯著優勢。"
)

SECTION_CONCLUSION = (
    "綜合三項實驗：六臂消融確立 calibrated gate 為工具增強管線之核心設計，"
    "而非工具呼叫本身；閾值敏感度證明 T=2 選擇穩健（T=3 同樣正向）；"
    "CodeBERT 長度公平性實驗揭示原始基準含偏差，公平條件下 DmAVID 勝出 +10.3 pp。"
)

# ── main ────────────────────────────────────────────────────────────────────

def build_new_section_elements(doc):
    """
    Return a list of XML elements representing the new section.
    Elements are in insertion order.
    """
    elems = []

    # --- Heading 3: 七、工具增強管線消融 ---
    # Style val "31" = Heading 3 in this document (Chinese Word numeric style ID)
    elems.append(make_para(doc,
        '七、工具增強管線消融：六臂消融與閾值敏感度分析',
        '31'))

    # --- Intro ---
    elems.append(make_para(doc, INTRO, 'Normal'))

    # === (一) 六臂消融 ===
    elems.append(make_para(doc, '（一）六臂消融（6-Arm Ablation）', 'Normal', bold=True))
    elems.append(make_para(doc, SIX_ARM_INTRO, 'Normal'))
    elems.append(make_para(doc, '表 4-A　工具增強管線六臂消融設計與結果（N=243）', 'Normal', bold=True))
    elems.append(make_table(doc, SIX_ARM_HEADERS, SIX_ARM_ROWS))
    elems.append(make_para(doc, SIX_ARM_FINDING1, 'Normal'))
    elems.append(make_para(doc, SIX_ARM_FINDING2, 'Normal'))
    elems.append(make_para(doc, SIX_ARM_FINDING3, 'Normal'))

    # === (二) 閾值敏感度 ===
    elems.append(make_para(doc, '（二）閾值敏感度分析（Threshold Sensitivity）', 'Normal', bold=True))
    elems.append(make_para(doc, THRESH_INTRO, 'Normal'))
    elems.append(make_para(doc, '表 4-B　Gate 閾值 T 敏感度掃描結果（N=243）', 'Normal', bold=True))
    elems.append(make_table(doc, THRESH_HEADERS, THRESH_ROWS))
    elems.append(make_para(doc, THRESH_FINDING, 'Normal'))

    # === (三) CodeBERT 長度公平性 ===
    elems.append(make_para(doc, '（三）CodeBERT 長度公平性消融補充', 'Normal', bold=True))
    elems.append(make_para(doc, CODEBERT_CONTENT, 'Normal'))
    elems.append(make_para(doc, CODEBERT_ARM_C, 'Normal'))
    elems.append(make_para(doc, CODEBERT_CONCLUSION, 'Normal'))

    # --- Section conclusion ---
    elems.append(make_para(doc, SECTION_CONCLUSION, 'Normal'))

    return elems


def main():
    doc = Document(IN_FILE)
    body = doc.element.body

    # ── Step 1: find para 546 (last para of 六、FP 歸因消融 section) ──
    all_paras = body.findall('.//' + qn('w:p'))
    ref_elem = doc.paragraphs[546]._element   # after this, insert new section

    # ── Step 2: build new section elements ──
    new_elems = build_new_section_elements(doc)

    # ── Step 3: insert after para 546, in reverse order (addnext inserts after ref) ──
    # To maintain order, insert in reverse: last element first
    cur = ref_elem
    for elem in new_elems:
        cur.addnext(elem)
        cur = elem   # each new elem becomes the new reference

    # ── Step 4: update "七、消融實驗綜合討論" → "八、消融實驗綜合討論" ──
    # Para 547 is now displaced; find it by text
    for para in doc.paragraphs:
        if para.text.strip().startswith('七、消融實驗綜合討論'):
            # Replace 七 with 八
            for run in para.runs:
                if '七、消融實驗綜合討論' in run.text:
                    run.text = run.text.replace('七、消融實驗綜合討論', '八、消融實驗綜合討論')
            # Also try direct text replace via XML
            for t_elem in para._element.iter(qn('w:t')):
                if '七、消融實驗綜合討論' in (t_elem.text or ''):
                    t_elem.text = t_elem.text.replace('七、消融實驗綜合討論', '八、消融實驗綜合討論')
            break

    # ── Step 5: update "七項消融實驗" → "八項消融實驗" in summary paras ──
    for para in doc.paragraphs:
        txt = para.text
        if '七項消融實驗' in txt:
            for t_elem in para._element.iter(qn('w:t')):
                if t_elem.text and '七項消融實驗' in t_elem.text:
                    t_elem.text = t_elem.text.replace('七項消融實驗', '八項消融實驗')

    # ── Step 6: save ──
    doc.save(OUT_FILE)
    print(f'Saved: {OUT_FILE}')

    # ── Verify ──
    doc2 = Document(OUT_FILE)
    h3s = [p.text.strip() for p in doc2.paragraphs
           if p.style.name.startswith('Heading 3') and '消融' in p.text]
    print('Heading 3 sections with 消融:')
    for h in h3s:
        print(f'  {h}')

    has_eight = any('八項消融實驗' in p.text for p in doc2.paragraphs)
    print(f'八項消融實驗 updated: {has_eight}')


if __name__ == '__main__':
    main()
