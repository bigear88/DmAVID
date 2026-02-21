#!/usr/bin/env python3
"""修改文獻探討 Word 文件 - 根據教授A和教授B的建議"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

INPUT_FILE = "/home/ubuntu/upload/基於大型語言模型的乙太坊DeFi智能合約漏洞偵測機制研究--文獻探討.docx"
OUTPUT_FILE = "/home/ubuntu/基於大型語言模型的乙太坊DeFi智能合約漏洞偵測機制研究--文獻探討_v2.docx"

doc = Document(INPUT_FILE)

# ============================================================
# 1. 修改表 2-2：將「本研究方案」的描述改為更謹慎的版本
# ============================================================
print("=" * 60)
print("步驟 1：修改表 2-2 中過於樂觀的描述")
print("=" * 60)

table_22_found = False
for table_idx, table in enumerate(doc.tables):
    # 找到表 2-2（包含「本研究的 ChatGPT + LLM 方案」的表格）
    for row_idx, row in enumerate(table.rows):
        row_text = " ".join([cell.text.strip() for cell in row.cells])
        if "本研究" in row_text and ("ChatGPT" in row_text or "LLM" in row_text):
            print(f"  找到表 2-2 第 {row_idx} 行：{row_text[:80]}...")
            table_22_found = True
            # 修改各欄位
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                # 找到並修改特定欄位
                # 根據文獻探討，表 2-2 的欄位順序：
                # 方法, 檢測準確率, 誤報率, 漏報率, 執行效率, 可擴展性, 技術門檻, 可解釋性, 適應性, 成本效益, 對DeFi特有漏洞的檢測能力
                pass
            break

# 由於表格結構複雜，改用段落文字替換方式
print("\n  嘗試在段落中找到表 2-2 相關描述...")

replacements_made = 0

for para_idx, para in enumerate(doc.paragraphs):
    original_text = para.text
    
    # 修改 1：表 2-2 後面的過於樂觀描述
    if "從表格中可以看出，本研究方案在檢測準確率、執行效率、可擴展性、可解釋性、對新型漏洞的適應性和對 DeFi 特有漏洞的檢測能力方面都具有優勢" in original_text:
        new_text = original_text.replace(
            "從表格中可以看出，本研究方案在檢測準確率、執行效率、可擴展性、可解釋性、對新型漏洞的適應性和對 DeFi 特有漏洞的檢測能力方面都具有優勢。然而，本方案的技術門檻和成本效益仍需要進一步的評估和優化。",
            "從表格中可以看出，本研究方案在可解釋性、對新型漏洞的適應性和對 DeFi 特有漏洞的檢測能力方面具有理論優勢。然而，上述評估係基於文獻分析與理論推導，實際性能表現（特別是誤報率與漏報率）仍需透過嚴謹的實驗驗證。值得注意的是，LLM 方法在實際應用中可能面臨較高的誤報率，此為後續第肆章實驗所需重點驗證之議題。本方案的技術門檻、API 呼叫成本及成本效益亦需要進一步的實證評估與優化。"
        )
        if new_text != original_text:
            # 保留格式，替換文字
            for run in para.runs:
                if "從表格中可以看出" in run.text:
                    run.text = run.text.replace(
                        "從表格中可以看出，本研究方案在檢測準確率、執行效率、可擴展性、可解釋性、對新型漏洞的適應性和對 DeFi 特有漏洞的檢測能力方面都具有優勢。然而，本方案的技術門檻和成本效益仍需要進一步的評估和優化。",
                        "從表格中可以看出，本研究方案在可解釋性、對新型漏洞的適應性和對 DeFi 特有漏洞的檢測能力方面具有理論優勢。然而，上述評估係基於文獻分析與理論推導，實際性能表現（特別是誤報率與漏報率）仍需透過嚴謹的實驗驗證。值得注意的是，LLM 方法在實際應用中可能面臨較高的誤報率，此為後續第肆章實驗所需重點驗證之議題。本方案的技術門檻、API 呼叫成本及成本效益亦需要進一步的實證評估與優化。"
                    )
            replacements_made += 1
            print(f"  ✓ 段落 {para_idx}：修改表 2-2 後的過於樂觀描述")

# ============================================================
# 2. 修改表 2-2 中「本研究方案」行的數值
# ============================================================
print("\n" + "=" * 60)
print("步驟 2：修改表 2-2 中本研究方案的評級")
print("=" * 60)

for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        row_text = " ".join([cell.text.strip() for cell in row.cells])
        if "本研究" in row_text and ("ChatGPT" in row_text or "LLM" in row_text):
            print(f"  表 {table_idx}, 行 {row_idx}")
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                print(f"    Cell {cell_idx}: '{cell_text}'")
            
            # 修改誤報率：「低」→「中（待驗證）」
            # 修改漏報率：「低」→「中（待驗證）」  
            # 修改檢測準確率：「高」→「高（待驗證）」
            # 需要根據實際的 cell 位置來修改
            for cell_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for run in p.runs:
                        # 修改特定欄位
                        pass
            break

# ============================================================
# 3. 修改 GPTScan/AuditGPT 比較段落
# ============================================================
print("\n" + "=" * 60)
print("步驟 3：修改 GPTScan/AuditGPT 比較說明")
print("=" * 60)

for para_idx, para in enumerate(doc.paragraphs):
    original_text = para.text
    
    # 找到提及 GPTScan 和 AuditGPT 作為比較基線的段落
    if "為了客觀評估本研究提出方案的有效性，實驗設計將引入兩個最新的、具代表性的 LLM 審計工具作為比較基線" in original_text:
        # 替換整段文字
        old_text = ("為了客觀評估本研究提出方案的有效性，實驗設計將引入兩個最新的、具代表"
                   "性的 LLM 審計工具作為比較基線：GPTScan（Sun et al., 2024）[20] 與 "
                   "AuditGPT（Xia et al., 2024）[21]。GPTScan 結合了 GPT 與靜態程序分析，"
                   "專注於檢測邏輯漏洞，是首個將 LLM 與程式分析相結合的智能合約漏洞檢測工具；"
                   "AuditGPT 則利用 ChatGPT 自動驗證智能合約是否符合 ERC 規則。透過與這些先"
                   "進工具進行效能比較，將能更清晰地定位本研究方案在檢測準確率、漏洞覆蓋範圍"
                   "及可解釋性等方面的相對優勢與獨特貢獻。")
        
        new_text = ("在相關研究中，GPTScan（Sun et al., 2024）[20] 與 AuditGPT（Xia et al., 2024）"
                   "[21] 是兩個具代表性的 LLM 審計工具。GPTScan 結合了 GPT 與靜態程序分析，"
                   "專注於檢測邏輯漏洞，是首個將 LLM 與程式分析相結合的智能合約漏洞檢測工具；"
                   "AuditGPT 則利用 ChatGPT 自動驗證智能合約是否符合 ERC 規則。然而，本研究"
                   "在實驗設計上未將此二工具納入直接比較基線，原因如下：（1）GPTScan 的原始碼"
                   "雖已公開，但其檢測範圍主要針對邏輯漏洞（Logic Vulnerabilities），與本研究涵蓋"
                   "重入攻擊、整數溢位、存取控制等多種漏洞類型的實驗範疇不完全對應；（2）AuditGPT "
                   "專注於 ERC 規則合規性驗證，其檢測目標與本研究的漏洞偵測目標存在本質差異；"
                   "（3）兩者所使用的評估數據集與本研究採用的 SmartBugs Curated 數據集不同，"
                   "直接比較可能產生不公平的結論。因此，本研究選擇以 Slither（靜態分析）和 "
                   "Mythril（符號執行）作為傳統工具基線，以單獨的 LLM（GPT-4.1-mini）作為 "
                   "LLM 基線，透過在相同數據集上的公平比較，來驗證本研究提出之混合式框架的有效性。"
                   "未來研究可在統一的評估基準下，進一步與 GPTScan、AuditGPT 等工具進行比較。")
        
        # 嘗試替換
        full_para_text = para.text
        if "為了客觀評估本研究提出方案的有效性" in full_para_text:
            # 清除所有 runs 並重新寫入
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            replacements_made += 1
            print(f"  ✓ 段落 {para_idx}：修改 GPTScan/AuditGPT 比較說明")

# ============================================================
# 4. 在第七節「研究缺口」之後補充研究限制
# ============================================================
print("\n" + "=" * 60)
print("步驟 4：補充研究限制說明")
print("=" * 60)

# 找到「本研究旨在填補上述研究缺口」的段落
target_para_idx = None
for para_idx, para in enumerate(doc.paragraphs):
    if "本研究旨在填補上述研究缺口" in para.text:
        target_para_idx = para_idx
        print(f"  找到目標段落 {para_idx}：{para.text[:60]}...")
        break

if target_para_idx is not None:
    # 在該段落之後插入研究限制段落
    # 由於 python-docx 不直接支援在指定位置插入段落，我們需要修改 XML
    from docx.oxml.ns import qn
    
    target_element = doc.paragraphs[target_para_idx]._element
    
    # 建立新段落
    limitation_texts = [
        "\n三、研究限制與預期挑戰",
        "本研究在設計與執行過程中，預期將面臨以下限制與挑戰，茲說明如下：",
        "（一）LLM 版本漂移風險：大型語言模型的 API 服務會定期更新模型版本，不同版本之間的行為可能存在差異。本研究將明確記錄實驗所使用的模型版本（如 GPT-4.1-mini）及實驗執行日期（2025 年 2 月），以確保實驗的可重現性。然而，未來使用者在不同時間點重現實驗時，可能因模型版本更新而獲得略有差異的結果。",
        "（二）資料集代表性限制：本研究主要採用 SmartBugs Curated 數據集作為實驗基準，該數據集雖為學術界廣泛使用的標準化數據集，但其合約樣本主要來自較早期的 Solidity 版本（0.4.x–0.5.x），可能無法完全反映當前 DeFi 生態系統中最新的合約模式與漏洞類型。此外，該數據集中的 DeFi 專屬漏洞（如閃電貸攻擊、預言機操縱）樣本數量有限，可能影響模型在此類漏洞上的評估充分性。",
        "（三）評估範圍限制：本研究的實驗評估主要聚焦於合約級別的二元分類（有漏洞/無漏洞），尚未深入評估函數級別或行級別的漏洞定位精確度。此外，LLM 生成的漏洞解釋與修復建議之品質評估，目前仍缺乏統一的量化標準。",
        "（四）成本與可擴展性：使用商業 LLM API（如 OpenAI GPT-4 系列）進行大規模合約審計會產生顯著的 API 呼叫成本。本研究將在實驗中記錄 Token 使用量與估算成本，以供實務應用時參考。",
        "（五）外部效度限制：本研究的實驗結果主要基於已知漏洞的合約樣本，對於零日漏洞（Zero-day Vulnerabilities）或全新的攻擊向量，模型的檢測能力仍有待進一步驗證。"
    ]
    
    for i, text in enumerate(limitation_texts):
        new_para = deepcopy(doc.paragraphs[target_para_idx]._element)
        # 清除內容
        for child in list(new_para):
            if child.tag.endswith('}r'):
                new_para.remove(child)
        
        # 建立新的 run
        new_run = deepcopy(doc.paragraphs[target_para_idx].runs[0]._element) if doc.paragraphs[target_para_idx].runs else None
        if new_run is not None:
            # 清除文字
            for t in new_run.findall(qn('w:t')):
                t.text = text
            new_para.append(new_run)
        
        target_element.addnext(new_para)
        target_element = new_para
    
    replacements_made += 1
    print(f"  ✓ 在段落 {target_para_idx} 之後插入研究限制（6 段）")

# ============================================================
# 5. 修改表 2-2 中本研究方案的具體欄位值
# ============================================================
print("\n" + "=" * 60)
print("步驟 5：修改表 2-2 中本研究方案的具體欄位值")
print("=" * 60)

for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        row_text = " ".join([cell.text.strip() for cell in row.cells])
        if "本研究" in row_text and ("ChatGPT" in row_text or "LLM" in row_text):
            print(f"  修改表 {table_idx}, 行 {row_idx}")
            for cell_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for run in p.runs:
                        # 誤報率：低 → 待驗證
                        if run.text.strip() == "低" and cell_idx in [1, 2, 3]:
                            # 根據欄位位置判斷
                            pass
                        # 漏報率：低 → 待驗證
            
            # 直接修改特定 cell 的文字
            cells = row.cells
            for cell_idx, cell in enumerate(cells):
                ct = cell.text.strip()
                if ct == "低" and cell_idx > 0:
                    # 判斷是否為誤報率或漏報率欄位
                    # 檢查表頭
                    header_row = table.rows[0]
                    header_text = header_row.cells[cell_idx].text.strip() if cell_idx < len(header_row.cells) else ""
                    if "誤報" in header_text or "漏報" in header_text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text.strip() == "低":
                                    run.text = "中（待驗證）"
                                    print(f"    ✓ Cell {cell_idx} ('{header_text}'): '低' → '中（待驗證）'")
                                    replacements_made += 1
            break

# ============================================================
# 6. 在第五節 LLM 方法限制部分補充 API 版本漂移說明
# ============================================================
print("\n" + "=" * 60)
print("步驟 6：補充 LLM API 版本漂移說明")
print("=" * 60)

for para_idx, para in enumerate(doc.paragraphs):
    if "計算成本與延遲：調用先進的 LLM API 會產生顯著的計算成本和網絡延遲" in para.text:
        # 在此段落的文字末尾補充
        for run in para.runs:
            if "計算成本與延遲" in run.text:
                run.text = run.text.rstrip()
                if not run.text.endswith("。"):
                    run.text += "。"
                run.text += "此外，商業 LLM API 的模型版本會定期更新，不同版本之間的行為差異可能影響實驗的可重現性，研究者應明確記錄所使用的模型版本與實驗日期。"
                replacements_made += 1
                print(f"  ✓ 段落 {para_idx}：補充 API 版本漂移說明")
                break

# ============================================================
# 保存
# ============================================================
print("\n" + "=" * 60)
print(f"保存修正後的文件...")
doc.save(OUTPUT_FILE)
import os
file_size = os.path.getsize(OUTPUT_FILE) / 1024
print(f"✓ 已保存至: {OUTPUT_FILE}")
print(f"  文件大小: {file_size:.1f} KB")
print(f"  共完成 {replacements_made} 項修改")
