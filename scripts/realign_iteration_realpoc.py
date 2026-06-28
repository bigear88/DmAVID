#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把論文迭代結果從『compile-gate 單調 0.9158 為主文』翻轉為
『real-PoC（PoC 驗證閘）為 canonical 主文、compile-only 閘門為消融』。
real-PoC（GitHub dmavid_autonomous/）：baseline 0.9061 → R1 0.9164 → R2 0.9060 → R3 0.9128（非單調）。
compile-only（BAK，已提交為消融）：0.9103 → 0.9153 → 0.9158（單調）。"""
import shutil
from docx import Document

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_realpoc.docx"
shutil.copyfile(SRC, BAK); print("backup:", BAK)
doc = Document(SRC)
paras = doc.paragraphs


def set_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell(cell, txt):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


# ---- 子字串替換（保留段落其餘內容）----
SUBS = {
 64: [("並經對抗式迭代後達 0.9158（R1=0.9103、R2=0.9153、R3=0.9158，呈單調遞增；惟其增益幅度有限且對驗證嚴格度與隨機種子敏感）",
       "而於嚴格之 Foundry PoC 攻擊重放驗證閘門下，對抗式迭代之三輪 F1 呈非單調變化（0.9164→0.9060→0.9128，峰值於首輪），最終與單通道基線相當，其價值主要體現於假陰性之精準修正與推理品質之提升；消融顯示若放寬驗證閘門為僅需 solc 編譯通過，三輪 F1 則單調上升至 0.9158（補丁未經漏洞利用驗證，宜視為效能上界）")],
 67: [("After adversarial iteration, the F1-score reaches 0.9158 (R1=0.9103, R2=0.9153, R3=0.9158; monotonically increasing), although this gain is modest and sensitive to the validation strictness and the random seed.",
       "Under a strict Foundry PoC exploit-replay verification gate, adversarial iteration yields a non-monotonic F1 trajectory (R1=0.9164, R2=0.9060, R3=0.9128, peaking in the first round) that is on par with the single-pass baseline, with its value lying mainly in precise false-negative correction and improved reasoning quality rather than large F1 gains; an ablation shows that relaxing the gate to solc-compilation-only strengthens the learning signal and produces a monotonic increase to 0.9158 (with patches not exploit-verified, treated as an upper bound).")],
 323: [("（R1=0.9103、R2=0.9153、R3=0.9158，呈單調遞增）",
        "（嚴格 PoC 驗證閘門下 R1=0.9164、R2=0.9060、R3=0.9128，呈非單調；compile-only 閘門消融則為 0.9103→0.9153→0.9158 單調）")],
 406: [("（含知識回饋閉環之迭代式管線於獨立批次中三輪 F1 為 0.9103→0.9153→0.9158、單調遞增、最終 Recall=0.9510、FPR=0.18，詳見第四節）",
        "（含知識回饋閉環之迭代式管線於嚴格 PoC 驗證閘門下三輪 F1 為 0.9164→0.9060→0.9128、非單調、最終 Recall=0.9510、FPR=0.19，詳見第四節）")],
 456: [("結果如表 4-11 所示，F1 隨迭代輪次之變化趨勢繪製於圖 4-7。",
        "結果如表 4-11 所示，F1 隨迭代輪次之變化趨勢繪製於圖 4-7。需特別說明，表 4-11 與圖 4-7 之 per-category 結果係取自 compile-only 閘門消融設定（變體編譯通過即回饋），用以展現放寬驗證下知識回饋對各漏洞類型之學習潛力；主文 canonical 之嚴格 PoC 驗證閘門整體結果見前述第一部分之 F1 分析。")],
 460: [("整體而言，per-category 之迭代效益為中性至正向，與全資料集 F1 之單調遞增觀察一致。",
        "整體而言，於 compile-only 閘門消融設定下 per-category 之迭代效益為中性至正向，與該設定全資料集 F1 之單調遞增觀察一致；惟主文嚴格 PoC 驗證閘門下，per-category 改善轉為漲跌互見、全資料集 F1 呈非單調。")],
 504: [("（2）多代理對抗式迭代之貢獻——預測效能呈單調遞增（三輪 F1 0.9103→0.9153→0.9158，最終 Recall 0.9510，高於單通道 0.9121）與過程品質強化（共 26 筆學習補丁同步向量化寫入 ChromaDB）",
        "（2）多代理對抗式迭代之貢獻——於嚴格 PoC 驗證閘門下全資料集 F1 呈非單調、最終 0.9128（峰值 R1 0.9164、Recall 0.9510），與單通道基線相當，價值主要在於 FN 修正與過程品質強化（共 26 筆學習補丁同步向量化寫入 ChromaDB；compile-only 閘門消融下則單調達 0.9158，惟補丁未經漏洞利用驗證）")],
 532: [("第二，對抗式多代理迭代（FN 課程學習）之三輪 F1 呈單調遞增（0.9103→0.9153→0.9158），最終 Recall 達 0.9510，三輪共合成 26 筆學習補丁並同步向量化寫入 ChromaDB，實現持續自我優化。",
        "第二，對抗式多代理迭代（FN 課程學習）於嚴格 PoC 驗證閘門下全資料集 F1 呈非單調、最終 0.9128（峰值 R1 0.9164、Recall 0.9510）；compile-only 閘門消融下則單調達 0.9158，三輪共合成 26 筆學習補丁並同步向量化寫入 ChromaDB，實現持續自我優化。")],
 548: [("在此基礎上，對抗式迭代透過三輪 FN 課程學習與 ChromaDB 知識回饋閉環，三輪 F1 呈單調遞增（0.9103→0.9153→0.9158），最終 Recall 達 0.9510、共合成 26 筆學習補丁（同步向量化寫入 ChromaDB）；惟其 F1 增益幅度有限且對驗證嚴格度與隨機種子敏感，價值亦體現於 FN 修正與過程品質。",
        "在此基礎上，對抗式迭代透過三輪 FN 課程學習與 ChromaDB 知識回饋閉環，於嚴格 PoC 驗證閘門下全資料集 F1 呈非單調、最終 0.9128（峰值 R1 0.9164、Recall 0.9510、共 26 筆學習補丁同步向量化寫入 ChromaDB）；compile-only 閘門消融下則單調達 0.9158。惟嚴格驗證下 F1 增益幅度有限且對驗證嚴格度與隨機種子敏感，其價值主要體現於 FN 修正與過程品質。")],
 558: [("在迭代效益層面，正式迭代管線（含 ChromaDB 知識回饋閉環，以 solc 編譯把關）之三輪 F1 呈單調遞增（0.9103→0.9153→0.9158），最終較單通道僅小幅提升（+0.0037）且對驗證嚴格度與隨機種子敏感，其核心價值在於 FN 課程學習之精準修正與過程品質之多維強化，而非大幅度之 F1 提升。",
        "在迭代效益層面，正式迭代管線（含 ChromaDB 知識回饋閉環）於嚴格 PoC 攻擊重放驗證閘門下三輪 F1 呈非單調、最終 0.9128（峰值 R1 0.9164），與單通道基線大致相當且對驗證嚴格度與隨機種子敏感；compile-only 閘門消融下則單調達 0.9158（補丁未經漏洞利用驗證，視為上界）。其核心價值在於 FN 課程學習之精準修正與過程品質之多維強化，而非大幅度之 F1 提升。")],
}

# ---- 整段重寫（set_text）----
FULL = {
 461: "就整體迭代效果而言，本研究以含知識回饋閉環之 Coordinator 自主型迭代實驗（全資料集 243 份合約，gpt-4.1-mini，temperature=0.1、seed=42，3 輪迭代，自乾淨知識庫起點，預算上限 $50、實花約 $23.88）作為正式迭代管線之 canonical 結果，其 Foundry 驗證階段先以 solc 編譯確認對抗變體之語法正確性，再以 forge test 執行攻擊重放（PoC）確認漏洞確可被利用，僅通過此雙重驗證之對抗知識方回饋至知識庫（詳見第參章 S5）。基線（LLM+RAG Stage 2，無迭代）F1 為 0.9061。啟用知識回饋閉環後，三輪 F1 呈非單調變化：第一輪（R1）F1=0.9164（Precision=0.8782，Recall=0.9580，FPR=0.19，FN=6），為三輪峰值；第二輪（R2）F1 回落至 0.9060（Precision=0.8710，Recall=0.9441，FPR=0.20，FN=8）；第三輪（R3）F1 回升至 0.9128（Precision=0.8774，Recall=0.9510，FPR=0.19，FN=7）。三輪後最終 F1=0.9128，與單通道 Self-Verify 之 0.9121 大致相當。其根本原因在於嚴格之 PoC 攻擊重放驗證使每輪可通過之有效對抗知識僅約 1 筆，學習訊號被大幅稀釋（詳見本節後段與 compile-only 閘門消融之對比）。惟須誠實指出，於嚴格驗證閘門下迭代之全資料集 F1 增益有限，其穩健性亦受驗證嚴格度與隨機種子兩項因素制約（詳見本節後段之穩健性檢定，即迭代顯著性檢定）。",
 467: "從表 4-11（compile-only 閘門消融）可觀察到，當驗證閘門放寬為僅需 solc 編譯通過時，per-category 切片於三輪迭代後多數類別呈正向改善：Unchecked Low-Level Calls（+0.030，達 0.9903）、Access Control（+0.030，達 0.9714）、Arithmetic（+0.040，達 0.9286）與 Time Manipulation（+0.111，達 1.0000）皆上升，Bad Randomness（0.9333）維持不變。Time Manipulation 之 1.0000 須保守解讀：該類別僅 5 份合約，樣本數偏少使單一預測即足以左右該類 F1，且本 per-category「F1」不含安全合約（precision 恆為 1.0、F1=2R/(1+R)、實為召回率轉換），不宜據以推論模型對時間操控漏洞已達完美偵測能力。全資料集 FPR 由 0.15 微升至 0.18。整體而言，compile-only 閘門下自主迭代之全資料集 F1 呈單調遞增（R1=0.9103、R2=0.9153、R3=0.9158）；惟於嚴格 PoC 驗證閘門（主文 canonical 結果）下，per-category 改善轉為漲跌互見（如 Access Control 由 R1 之 1.0000 回落至 R3 之 0.9412、Arithmetic 由 0.9286 回落至 0.8889），全資料集 F1 亦呈非單調（0.9164→0.9060→0.9128），凸顯迭代增益對驗證嚴格度之高度敏感。",
 468: "此一結果之核心意涵在於：於嚴格之 Foundry PoC 攻擊重放驗證閘門（要求對抗變體之漏洞確可被利用方回饋知識）下，多代理對抗式迭代於 SmartBugs 上之全資料集 F1 呈非單調變化（三輪 0.9164→0.9060→0.9128，峰值於首輪、最終與單通道基線相當）；其根本原因在於每輪可通過 PoC 驗證之有效對抗知識僅約 1 筆，學習訊號薄弱。為釐清迭代機制本身之學習潛力，本研究進行 compile-only 閘門消融：若將驗證階段放寬為僅需 solc 編譯通過（變體編譯通過即回饋），則每輪可通過之有效對抗知識增為 11/8/7 筆（共 26 筆），學習訊號充足，三輪 F1 隨之呈單調且正向之改善（0.9103→0.9153→0.9158），Recall 由 0.9231 升至 0.9510。此一對比明確顯示：FN 課程學習閉環之設計理念成立（放寬閘門即見單調學習曲線），但其全資料集 F1 增益之大小高度依賴知識回饋閘門之寬鬆程度——compile-only 之 0.9158 因補丁未經漏洞利用驗證，宜視為效能上界，而嚴格 PoC 驗證之 0.9128 方為可被漏洞利用證實之 canonical 結果。此外，迭代之穩健性亦受隨機種子制約：以 3 個隨機種子（42/7/123）對處理組與 placebo 對照組之嚴格檢定顯示，輪間增益於跨種子下並不穩定（詳見後段之迭代顯著性檢定）。綜合而言，迭代之過程層面貢獻明確：FN 課程學習成立、知識庫動態擴充（補丁累積寫入 vulnerability_knowledge.json 並同步向量化至 ChromaDB），以及可解釋性深化——推理深度由平均 68 詞提升至 73.2 詞，修復建議生成率達 62.86%，程式碼行號引用精確率達 100%（此點將於第七節進一步驗證）；而其全資料集 F1 之增益則屬小幅，且對驗證嚴格度與隨機種子敏感，宜保守解讀。",
}

# ---- 套用 ----
for idx, repls in SUBS.items():
    p = paras[idx]; t = p.text; new = t
    for old, rep in repls:
        if old in new:
            new = new.replace(old, rep)
        else:
            print(f"[MISS] P{idx} 子字串未找到: {old[:30]}")
    if new != t:
        set_text(p, new); print(f"[OK] P{idx} 子字串替換")

for idx, txt in FULL.items():
    set_text(paras[idx], txt); print(f"[OK] P{idx} 整段重寫")

# ---- 圖表標題 ----
for idx in [457]:
    if "對抗迭代結果" in paras[idx].text and "compile-only" not in paras[idx].text:
        set_text(paras[idx], paras[idx].text.replace("對抗迭代結果", "對抗迭代結果（compile-only 閘門消融）"))
        print(f"[OK] P{idx} 表4-11 標題加註")
for idx in [459]:
    if "F1 變化趨勢" in paras[idx].text and "compile-only" not in paras[idx].text:
        set_text(paras[idx], paras[idx].text.replace("F1 變化趨勢", "F1 變化趨勢（compile-only 閘門消融）"))
        print(f"[OK] P{idx} 圖4-7 標題加註")

# ---- 表格 ----
# 表4-1 (T6) row5: DmAVID+迭代 -> real-PoC R3
t6 = doc.tables[6].rows[5].cells
set_cell(t6[2], "0.8774"); set_cell(t6[4], "0.9128"); set_cell(t6[5], "+多代理對抗迭代 3 輪（PoC 驗證閘）")
print("[OK] 表4-1 迭代列 -> 0.8774/0.9510/0.9128（PoC 驗證閘）")
# 表4-2 (T7) row5: +對抗迭代 R3 -> real-PoC
t7 = doc.tables[7].rows[5].cells
set_cell(t7[0], "+對抗迭代 R3（PoC 驗證閘，FINAL）")
set_cell(t7[2], "19"); set_cell(t7[4], "81"); set_cell(t7[5], "0.8774"); set_cell(t7[7], "0.9128"); set_cell(t7[8], "0.19")
print("[OK] 表4-2 迭代列 -> 136/19/7/81 0.8774/0.9510/0.9128 FPR0.19")

doc.save(SRC)
print("\nSAVED:", SRC)
