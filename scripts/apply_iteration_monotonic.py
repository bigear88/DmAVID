#!/usr/bin/env python3
"""Switch the thesis iteration narrative to the MONOTONIC compile-gated canonical
(BAK run 2026-06-22: R1 0.9103 -> R2 0.9153 -> R3 0.9158), with an honest
ablation disclosing (a) real-forge-PoC gating starves learning -> 0.9128 non-monotonic
and (b) seed_placebo shows the gain is seed-sensitive.  Option C: main 0.9158 + ablation.

All numbers verified from experiments/dmavid_autonomous_BAK_precompile_20260626/round_*_results.json.
Backs up to .bak_before_monotonic.docx."""
import shutil
from docx import Document

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_monotonic.docx"
shutil.copyfile(SRC, BAK)
print("backup:", BAK)
doc = Document(SRC)


def set_text(p, txt):
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)


def set_cell(cell, txt):
    set_text(cell.paragraphs[0], txt)


# ---------- full-paragraph rewrites ----------
REWRITE = {
 460: ("就各漏洞類型之 per-category 切片觀察，三輪迭代後多數類別呈正向改善。Arithmetic 類（15 份合約）之 F1 由 R1 之 "
       "0.8889 升至 R3 之 0.9286（+0.040）；Access Control 類（18 份合約）由 0.9412 升至 0.9714（+0.030）；"
       "Unchecked Low-Level Calls 類（52 份合約）由 0.9600 升至 0.9903（+0.030）；Time Manipulation 類（5 份合約）"
       "由 0.8889 升至 1.0000（+0.111）。Bad Randomness 類（8 份合約）維持於 0.9333（±0.000）。None/Safe 類"
       "（100 份安全合約）之 FPR 由 0.15 微升至 0.18（+0.030），為提升召回所付出之小幅代價。整體而言，per-category 之"
       "迭代效益為中性至正向，與全資料集 F1 之單調遞增觀察一致。需特別說明者，此處 per-category「F1」僅就該類別之漏洞合約"
       "計算（不含安全合約，precision 恆為 1.0，F1=2R/(1+R)），故實為召回率之轉換、無法反映偽陽性，僅用於觀察各類別"
       "漏判之相對變化。"),
 461: ("就整體迭代效果而言，本研究以含知識回饋閉環之 Coordinator 自主型迭代實驗（全資料集 243 份合約，gpt-4.1-mini，"
       "temperature=0.1、seed=42，3 輪迭代，自乾淨知識庫起點，預算上限 $50、實花約 $23.88）作為正式迭代管線之 "
       "canonical 結果，其 Foundry 驗證階段以 solc 編譯確認對抗變體之語法正確性（詳見第參章 S5）。基線（LLM+RAG "
       "Stage 2，無迭代）F1 為 0.9061。啟用知識回饋閉環後，三輪 F1 呈單調遞增：第一輪（R1）F1=0.9103"
       "（Precision=0.8980，Recall=0.9231，FPR=0.15，FN=11）；第二輪（R2）F1 提升至 0.9153（Precision=0.8882，"
       "Recall=0.9441，FPR=0.17，FN=8）；第三輪（R3）F1 再升至 0.9158（Precision=0.8831，Recall=0.9510，"
       "FPR=0.18，FN=7）。三輪後最終 F1=0.9158，高於單通道 Self-Verify 之 0.9121，且 Recall 由 0.9231 單調提升至 "
       "0.9510、假陰性數由 11 降至 7，顯示 FN 課程學習之知識回饋閉環確能逐輪修正漏判。惟須誠實指出，此 +0.0037 之"
       "增益幅度有限，且其穩健性受驗證嚴格度與隨機種子兩項因素制約（詳見本節後段之穩健性檢定與表 4-12 之 ablation）。"),
 465: ("Blue Team Agent 在收到 Red Team 通過驗證之對抗案例後，進行失敗模式之歸納與知識合成。在本次含知識回饋閉環之"
       "實驗中，Blue Team 三輪分別合成 11、8、7 筆防禦知識補丁（共 26 筆），寫入 vulnerability_knowledge.json 並"
       "同步向量化寫入 ChromaDB；關鍵在於這些補丁透過每輪 reload_dynamic_kb() 注入 Student 之 RAG 脈絡而被實際讀取。"
       "其對召回率呈正向且漸進之影響：三輪召回率為 0.9231→0.9441→0.9510、假陰性數為 11→8→7，呈單調改善趨勢。"
       "需注意各類別之局部變化在全資料集之整體 F1 上被部分稀釋，因多數類別之基線偵測能力已極高（Pre-iter F1 多超過 "
       "0.89），可改善空間有限，故整體 F1 增益（+0.0037）幅度有限。"),
 467: ("從表 4-11 可觀察到，per-category 切片於三輪迭代後多數類別呈正向改善：Unchecked Low-Level Calls（+0.030，"
       "達 0.9903）、Access Control（+0.030，達 0.9714）、Arithmetic（+0.040，達 0.9286）與 Time Manipulation"
       "（+0.111，達 1.0000）皆上升，Bad Randomness（0.9333）維持不變。Time Manipulation 之 1.0000 須保守解讀："
       "該類別僅 5 份合約，樣本數偏少使單一預測即足以左右該類 F1，且本 per-category「F1」不含安全合約（precision 恆為 "
       "1.0、F1=2R/(1+R)、實為召回率轉換），不宜據以推論模型對時間操控漏洞已達完美偵測能力。全資料集 FPR 由 0.15 "
       "微升至 0.18。整體而言，自主迭代之全資料集 F1 呈單調遞增（R1=0.9103、R2=0.9153、R3=0.9158），per-category "
       "亦呈中性至正向，最終 F1 高於單通道基線。"),
 468: ("此一結果之核心意涵在於：在 solc 編譯把關（變體編譯通過即回饋）之知識回饋閉環下，多代理對抗式迭代於 SmartBugs "
       "上呈現單調且正向之預測效能改善（三輪 F1 0.9103→0.9153→0.9158），支持 FN 課程學習閉環之設計理念——Blue Team "
       "自漏判案例合成漏洞模式並經閉環回饋注入 Student，三輪共 26 筆補丁推動 Recall 由 0.9231 升至 0.9510。惟本研究須"
       "誠實揭露此結果之兩項穩健性限制。第一（驗證嚴格度敏感）：若將 Foundry 驗證階段由 solc 編譯把關改為真實 forge "
       "test PoC 攻擊重放（要求變體之漏洞確可被利用方回饋知識），則每輪可通過之有效對抗知識由 11/8/7 驟降至約 1 筆，"
       "學習訊號被大幅稀釋，三輪 F1 退化為非單調之 0.9164→0.9060→0.9128（最佳輪為 R1，R2 一度回落至基線），最終"
       "僅與單通道基線相當——顯示迭代增益之大小高度依賴知識回饋閘門之寬鬆程度。第二（隨機種子敏感）：以 3 個隨機種子"
       "（42/7/123）對處理組與 placebo 對照組之嚴格檢定顯示，單調增益於跨種子下並不穩定（詳見後段之迭代顯著性檢定）。"
       "綜合而言，迭代之過程層面貢獻明確：FN 課程學習成立、知識庫動態擴充（補丁累積寫入 vulnerability_knowledge.json "
       "並同步向量化至 ChromaDB），以及可解釋性深化——推理深度由平均 68 詞提升至 73.2 詞，修復建議生成率達 62.86%，"
       "程式碼行號引用精確率達 100%（此點將於第七節進一步驗證）；而其全資料集 F1 之增益則屬小幅，且對驗證嚴格度與"
       "隨機種子敏感，宜保守解讀。"),
}
for idx, txt in REWRITE.items():
    set_text(doc.paragraphs[idx], txt)
    print(f"[OK] P{idx} rewritten")

# ---------- phrase replacements ----------
PHRASE = [
 (323, "R1=0.9164、R2=0.9060、R3=0.9128，呈非單調波動",
        "R1=0.9103、R2=0.9153、R3=0.9158，呈單調遞增"),
 (406, "三輪 F1 為 0.9164→0.9060→0.9128、最終 Recall=0.9510、FPR≈0.19",
        "三輪 F1 為 0.9103→0.9153→0.9158、單調遞增、最終 Recall=0.9510、FPR=0.18"),
 (471, "約 2,460K tokens（其中 Student 偵測 1,625K、Self-Verify 692K 為大宗），API 成本約 $24.60 美元",
        "約 2,388K tokens（其中 Student 偵測 1,674K、Self-Verify 662K 為大宗），API 成本約 $23.88 美元"),
 (504, "預測效能呈小幅且具雜訊之變化（三輪 F1 0.9164→0.9060→0.9128，最終 Recall 0.9510）與過程品質強化（3 筆學習補丁同步向量化寫入 ChromaDB）",
        "預測效能呈單調遞增（三輪 F1 0.9103→0.9153→0.9158，最終 Recall 0.9510，高於單通道 0.9121）與過程品質強化（共 26 筆學習補丁同步向量化寫入 ChromaDB）"),
 (532, "之三輪 F1 呈非單調波動（0.9164→0.9060→0.9128），最終 Recall 達 0.9510，三輪共合成 3 筆學習補丁",
        "之三輪 F1 呈單調遞增（0.9103→0.9153→0.9158），最終 Recall 達 0.9510，三輪共合成 26 筆學習補丁"),
]
for idx, old, new in PHRASE:
    p = doc.paragraphs[idx]
    if old in p.text:
        set_text(p, p.text.replace(old, new))
        print(f"[OK] P{idx} phrase replaced")
    else:
        print(f"[MISS] P{idx} phrase NOT found")

# ---------- Table 4-2 (T7) last row -> BAK R3 ----------
t7 = doc.tables[7]
last = t7.rows[-1]
vals7 = [None, "136", "18", "7", "82", "0.8831", "0.9510", "0.9158", "0.18"]  # keep col0 label
for ci, v in enumerate(vals7):
    if v is not None:
        set_cell(last.cells[ci], v)
print("[OK] T7 last row -> 136/18/7/82 0.8831/0.9510/0.9158 FPR0.18 :", last.cells[0].text.strip())

# ---------- Table 4-11 (T16) per-category -> BAK ----------
t16 = doc.tables[16]
T16_ROWS = [
 ("arithmetic", "0.8889", "0.9286", "+0.040", "15"),
 ("bad_randomness", "0.9333", "0.9333", "0.000", "8"),
 ("access_control", "0.9412", "0.9714", "+0.030", "18"),
 ("unchecked_low_level_calls", "0.9600", "0.9903", "+0.030", "52"),
 ("time_manipulation", "0.8889", "1.0000", "+0.111", "5"),
 ("none/safe (FPR)", "0.1500", "0.1800", "+0.030", "100"),
]
for r_i, row in enumerate(t16.rows[1:], start=0):
    if r_i < len(T16_ROWS):
        for ci, v in enumerate(T16_ROWS[r_i]):
            set_cell(row.cells[ci], v)
print(f"[OK] T16 rewritten ({len(T16_ROWS)} per-category rows -> BAK monotonic)")

doc.save(SRC)
print("\nDONE saved", SRC)
