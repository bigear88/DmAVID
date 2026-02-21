#!/usr/bin/env python3
"""
在論文正文 v5 中插入補充內容：混淆矩陣、McNemar 檢驗、成本敏感分析等
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import json
import os

INPUT_FILE = "/home/ubuntu/基於大型語言模型的乙太坊DeFi智能合約漏洞偵測機制研究--論文正文_v5.docx"
OUTPUT_FILE = "/home/ubuntu/基於大型語言模型的乙太坊DeFi智能合約漏洞偵測機制研究--論文正文_v5_final.docx"

doc = Document(INPUT_FILE)

# 載入補充數據
SUPP_DIR = "/home/ubuntu/defi-vuln-detection/supplementary_results"
with open(os.path.join(SUPP_DIR, "confusion_matrices.json")) as f:
    cm_data = json.load(f)
with open(os.path.join(SUPP_DIR, "mcnemar_tests.json")) as f:
    mcnemar_data = json.load(f)
with open(os.path.join(SUPP_DIR, "cost_sensitive_analysis.json")) as f:
    cost_data = json.load(f)

# ===== 工具函數 =====
def set_cell_text(cell, text, bold=False, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """設定表格儲存格的文字和格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    # 設定中文字體
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '標楷體'})
    rPr.append(rFonts)

def add_heading_paragraph(doc, element_before, text, level=3):
    """在指定元素後添加標題段落"""
    from docx.oxml import OxmlElement
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    element_before.addnext(p)
    return p

def add_body_paragraph(doc, element_before, text, font_size=12, bold=False, indent=True):
    """在指定元素後添加正文段落"""
    from docx.oxml import OxmlElement
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')  # 首行縮排
        pPr.append(ind)
    # 行距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')  # 1.5 倍行距
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b_elem = OxmlElement('w:b')
        rPr.append(b_elem)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '標楷體')
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    element_before.addnext(p)
    return p

def create_table_after(doc, element_before, headers, rows, caption=None):
    """在指定元素後創建表格"""
    from docx.oxml import OxmlElement
    
    # 先添加表格標題
    if caption:
        cap_p = add_body_paragraph(doc, element_before, caption, font_size=10, bold=True, indent=False)
        cap_pPr = cap_p.find(qn('w:pPr'))
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        cap_pPr.append(jc)
        element_before = cap_p
    
    ncols = len(headers)
    nrows = len(rows) + 1  # +1 for header
    
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    # 表格邊框
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)
    
    # 表格網格
    tblGrid = OxmlElement('w:tblGrid')
    col_width = 9000 // ncols
    for _ in range(ncols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(col_width))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)
    
    def make_row(cells, is_header=False):
        tr = OxmlElement('w:tr')
        for cell_text in cells:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            if is_header:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9E2F3')
                tcPr.append(shd)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '276')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')  # 9pt
            rPr.append(sz)
            if is_header:
                b = OxmlElement('w:b')
                rPr.append(b)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), '標楷體')
            rPr.append(rFonts)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(cell_text)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    
    # Header row
    make_row(headers, is_header=True)
    # Data rows
    for row in rows:
        make_row(row)
    
    element_before.addnext(tbl)
    return tbl

# ===== 找到插入位置 =====
print("尋找插入位置...")

# 找到「第四節 DeFi 專屬漏洞檢測實驗」之前的位置
target_idx = -1
for i, para in enumerate(doc.paragraphs):
    if "第四節 DeFi 專屬漏洞檢測實驗" in para.text:
        target_idx = i
        print(f"  找到「第四節 DeFi 專屬漏洞檢測實驗」在段落 {i}")
        break

if target_idx == -1:
    # 嘗試找「DeFi 專屬漏洞」
    for i, para in enumerate(doc.paragraphs):
        if "DeFi 專屬漏洞" in para.text:
            target_idx = i
            print(f"  找到「DeFi 專屬漏洞」在段落 {i}")
            break

if target_idx == -1:
    print("  ✗ 未找到插入位置，嘗試在第六節綜合討論之前插入")
    for i, para in enumerate(doc.paragraphs):
        if "第六節 綜合討論" in para.text:
            target_idx = i
            print(f"  找到「第六節 綜合討論」在段落 {i}")
            break

# 在目標段落之前插入（即在前一段之後）
insert_element = doc.paragraphs[target_idx - 1]._element

# ===== 插入內容（從後往前插入，因為 addnext 是在元素後面插入） =====
print("\n插入補充內容...")

# 7. 各漏洞類型檢測表現（最後插入，所以最先寫）
vuln_type_rows = [
    ["reentrancy", "31", "94%", "100%", "100%", "100%"],
    ["unchecked_low_level", "52", "96%", "100%", "98%", "100%"],
    ["access_control", "18", "94%", "100%", "100%", "100%"],
    ["arithmetic", "15", "93%", "93%", "87%", "87%"],
    ["bad_randomness", "8", "100%", "100%", "100%", "100%"],
    ["denial_of_service", "6", "83%", "100%", "100%", "100%"],
    ["front_running", "4", "100%", "100%", "100%", "100%"],
    ["time_manipulation", "5", "100%", "100%", "100%", "100%"],
]

# 先插入空行分隔
sep7 = add_body_paragraph(doc, insert_element, "", font_size=12, indent=False)

# 漏洞類型說明文字
vuln_text = "LLM 方法在大多數漏洞類型上都達到了 100% 的召回率，尤其在 reentrancy（重入攻擊）和 access_control（存取控制）等高風險漏洞上表現優異。Slither 在 arithmetic（整數溢位）類型上的表現與 LLM 方法相當，但在 denial_of_service 等需要語義理解的漏洞類型上略遜一籌。"
vuln_desc = add_body_paragraph(doc, sep7, vuln_text, font_size=12)

# 漏洞類型表格
vuln_tbl = create_table_after(doc, vuln_desc,
    ["漏洞類型", "合約數", "Slither", "LLM Base", "LLM+RAG", "Hybrid"],
    vuln_type_rows,
    caption="表 4-7：各漏洞類型的檢測召回率"
)

# 漏洞類型標題
vuln_heading = add_body_paragraph(doc, vuln_tbl, "七、各漏洞類型檢測表現比較", font_size=14, bold=True, indent=False)

print("  ✓ 插入漏洞類型比較表")

# 6. 成本敏感分析
sep6 = add_body_paragraph(doc, vuln_heading, "", font_size=12, indent=False)

cost_text = "分析結果顯示，在絕大多數成本比率（1x 至 20x）下，LLM+RAG 都是總成本最低的方法。僅在極端情況下（FN 成本為 FP 的 50 倍以上），純 LLM Base 因其近乎完美的召回率（99.30%）而成為最佳選擇，但代價是 98% 的誤報率。這一分析為實務應用提供了明確的方法選擇指引。"
cost_desc = add_body_paragraph(doc, sep6, cost_text, font_size=12)

cost_rows = [
    ["1x", "94", "99", "36", "56", "LLM+RAG"],
    ["2x", "101", "100", "39", "58", "LLM+RAG"],
    ["5x", "122", "103", "48", "64", "LLM+RAG"],
    ["10x", "157", "108", "63", "74", "LLM+RAG"],
    ["20x", "227", "118", "93", "94", "LLM+RAG"],
    ["50x", "437", "148", "183", "154", "LLM Base"],
]

cost_tbl = create_table_after(doc, cost_desc,
    ["FN/FP比率", "Slither", "LLM Base", "LLM+RAG", "Hybrid", "最佳方法"],
    cost_rows,
    caption="表 4-6：成本敏感分析（不同 FN/FP 成本比率下的總成本）"
)

cost_intro = add_body_paragraph(doc, cost_tbl, "在實際應用中，漏報（FN）與誤報（FP）的代價往往不對稱——漏掉一個真實漏洞可能導致數百萬美元的損失，而誤報僅增加人工審查的成本。表 4-6 呈現了在不同 FN/FP 成本比率下各方法的總成本比較。", font_size=12)

cost_heading = add_body_paragraph(doc, cost_intro, "六、成本敏感分析", font_size=14, bold=True, indent=False)

print("  ✓ 插入成本敏感分析")

# 5. McNemar 統計檢驗
sep5 = add_body_paragraph(doc, cost_heading, "", font_size=12, indent=False)

mcnemar_text = "結果顯示：(1) LLM+RAG 相較於純 LLM Base 的改善具有高度統計顯著性（p < 0.001），證實 RAG 知識庫的引入確實帶來了實質性的性能提升；(2) LLM+RAG 與 Hybrid 之間也存在顯著差異（p = 0.002），LLM+RAG 在整體準確率上優於 Hybrid；(3) 在僅包含漏洞合約的子集上，LLM+RAG 與 Slither 的差異未達統計顯著水準（p = 0.343），這是因為兩者在漏洞合約上的召回率都很高（>95%），差異主要體現在安全合約的誤報率上。"
mcnemar_desc = add_body_paragraph(doc, sep5, mcnemar_text, font_size=12)

mcnemar_rows = [
    ["LLM+RAG vs LLM Base", "243", "57.37", "<0.0001", "***"],
    ["LLM+RAG vs Hybrid", "243", "9.50", "0.0021", "***"],
    ["Hybrid vs LLM Base", "243", "37.53", "<0.0001", "***"],
    ["LLM+RAG vs Slither", "143", "0.90", "0.3428", "n.s."],
    ["Hybrid vs Slither", "143", "1.78", "0.1824", "n.s."],
]

mcnemar_tbl = create_table_after(doc, mcnemar_desc,
    ["比較對", "共同合約", "χ²", "p-value", "顯著性"],
    mcnemar_rows,
    caption="表 4-5：McNemar 統計檢驗結果（*** p < 0.01; n.s. = not significant）"
)

mcnemar_intro = add_body_paragraph(doc, mcnemar_tbl, "為了驗證各方法之間的性能差異是否具有統計顯著性，我們使用 McNemar 檢驗對配對預測結果進行分析。McNemar 檢驗適用於比較兩個分類器在相同測試集上的表現差異，其原假設為兩個分類器的錯誤率相同。表 4-5 列出了主要方法對之間的 McNemar 檢驗結果。", font_size=12)

mcnemar_heading = add_body_paragraph(doc, mcnemar_intro, "五、統計顯著性檢驗", font_size=14, bold=True, indent=False)

print("  ✓ 插入 McNemar 統計檢驗")

# 4. 混淆矩陣
sep4 = add_body_paragraph(doc, mcnemar_heading, "", font_size=12, indent=False)

cm_text = "從混淆矩陣中可以清楚看出各方法的特性：Slither 和純 LLM Base 都有極高的 TP 但也有極高的 FP（分別為 87 和 98），顯示它們傾向於將大部分合約標記為有漏洞。Mythril 則完全相反，FP 為 0 但 FN 高達 11（在 20 個漏洞合約中漏掉了 11 個）。LLM+RAG 在維持高 TP（140）的同時，將 FP 控制在 33，是所有方法中最佳的平衡點。"
cm_desc = add_body_paragraph(doc, sep4, cm_text, font_size=12)

cm_rows = [
    ["Slither", "136", "87", "7", "13", "0.6132", "0.6099", "0.9510", "0.7432", "0.8700"],
    ["Mythril*", "9", "0", "11", "20", "0.7250", "1.0000", "0.4500", "0.6207", "0.0000"],
    ["LLM Base", "142", "98", "1", "2", "0.5926", "0.5917", "0.9930", "0.7415", "0.9800"],
    ["LLM+RAG", "140", "33", "3", "67", "0.8519", "0.8092", "0.9790", "0.8861", "0.3300"],
    ["Hybrid", "141", "54", "2", "46", "0.7695", "0.7231", "0.9860", "0.8343", "0.5400"],
]

cm_tbl = create_table_after(doc, cm_desc,
    ["方法", "TP", "FP", "FN", "TN", "Acc", "Prec", "Rec", "F1", "FPR"],
    cm_rows,
    caption="表 4-4：各方法在 SmartBugs 測試集上的混淆矩陣"
)

cm_note = add_body_paragraph(doc, cm_tbl, "*Mythril 僅分析 40 個合約（20 漏洞 + 20 安全），因符號執行的計算成本極高（平均 36.2 秒/合約，47.5% 超時率）。", font_size=10, indent=False)

cm_intro = add_body_paragraph(doc, cm_note, "為了更清楚地呈現各方法的分類表現，表 4-4 列出了所有方法在 SmartBugs 測試集上的完整混淆矩陣與衍生指標。", font_size=12)

cm_heading = add_body_paragraph(doc, cm_intro, "四、混淆矩陣分析", font_size=14, bold=True, indent=False)

print("  ✓ 插入混淆矩陣")

# ===== 保存 =====
print("\n保存修正後的文件...")
doc.save(OUTPUT_FILE)
print(f"✓ 已保存至: {OUTPUT_FILE}")
print(f"  文件大小: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")
