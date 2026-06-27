#!/usr/bin/env python3
"""Apply the agreed thesis revisions (seed+placebo result) to the docx.
Backs up first; surgical phrase replacement preserving paragraph style."""
import shutil, sys
from docx import Document

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_seedplacebo.docx"
shutil.copyfile(SRC, BAK)
print("backup:", BAK)
doc = Document(SRC)

def set_text(p, txt):
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)

# (paragraph_index, old_phrase, new_phrase)
EDITS = [
 # C — soften design-intent claims
 (311,
  "賦予 DmAVID 突破單次推論性能瓶頸的能力，實現偵測模型與領域知識的持續演進。",
  "其設計目標為突破單次推論之性能瓶頸；惟此自我改進機制之實際效益，於第四節以 multi-seed 與 placebo 對照實驗嚴格檢定。"),
 (327,
  "透過對抗式迭代機制，DmAVID 系統具備持續自強化（self-strengthening）之能力，每一輪迭代均可能發現新的偵測盲區並予以修補。",
  "透過對抗式迭代機制，DmAVID 系統於設計上具備自強化（self-strengthening）機制，期能於每一輪迭代發現並修補偵測盲區；惟此機制之實證效益須以對照實驗檢定（詳見第四節）。"),
 (345,
  "因此其偵測能力隨迭代輪次之推進而逐步增強。",
  "此一機制意圖使偵測能力隨迭代輪次精進；惟第四節之 multi-seed 與 placebo 對照顯示，此增益落於測量雜訊之內。"),
 (326,
  "雙方之能力在此過程中共同提升。",
  "雙方之能力在此過程中共同提升（本研究於第四節以對照實驗檢驗此類比於本任務之實際成立程度）。"),
 # Seed reproducibility correction
 (383,
  "其次，所有隨機過程均以 seed=42 固定，保障實驗可重複性。",
  "其次，本地隨機過程（資料切分、CodeBERT 訓練、傳統 ML）以 seed=42 固定以確保可重複性；惟核心 LLM 代理經 OpenAI API 推論，其 seed 參數僅為盡力而為（best-effort），於 temperature=0.1 下輸出仍具非決定性（實測同設定兩次 F1 差異約 0.01），為嚴謹量化此變異，迭代效益評估另採 multi-seed（42、7、123）並輔以 placebo 對照（詳見第四節）。"),
 # CTC dimension correction
 (385,
  "在可解釋性方面，本研究採用正確性-完整性-清晰度（CTC, Correctness-Thoroughness-Clarity）指標體系，包含漏洞模式覆蓋率（Pattern Coverage）、根因準確度（Root Cause Accuracy）、攻擊路徑覆蓋率（Attack Path Coverage）及修復建議品質（Repair Quality）四個維度。",
  "在可解釋性方面，本研究採用正確性-完整性-清晰度（CTC, Correctness-Thoroughness-Clarity）三維度框架（Smart-LLaMA-DPO, ISSTA 2025），由 LLM 評審就正確性、完整性、清晰度各以 1–10 分評分，綜合分數依加權 CTC＝0.6×Correctness＋0.3×Thoroughness＋0.1×Clarity 計算；另以四項自動化可解釋性子指標（漏洞模式覆蓋率 Pattern Coverage、根因準確度 Root Cause Accuracy、攻擊路徑覆蓋率 Attack Path Coverage、修復建議品質 Repair Quality）作為細項分析（詳見第四節表 4-19）。"),
 (521,
  "如表 4-19 所示，DmAVID 之 CTC 各指標細項揭示了管線在可解釋性上之優勢與瓶頸。",
  "如表 4-19 所示，DmAVID 之四項自動化可解釋性子指標揭示了管線在可解釋性上之優勢與瓶頸。"),
 # D — reframe +0.0067 into placebo-controlled verdict
 (461,
  "三輪後最終 F1=0.9128，較 LLM+RAG 基線提升 +0.0067，淨效益為小幅正向但具明顯輪間雜訊——R1 即已捕獲主要增益，R2 一度回落至基線水準。",
  "三輪後最終 F1=0.9128；惟此單一執行之 +0.0067 落於測量雜訊內——經 multi-seed（42/7/123）與 placebo 對照之嚴格檢定（處理組 R3=0.9149±0.0109、placebo 組 R3=0.9201±0.0072，信賴帶重疊且 placebo 平均略高），確認此增益未達統計顯著（詳見本節後段之迭代顯著性檢定）。"),
 (468,
  "最終較 LLM+RAG 基線（0.9061）提升 +0.0067。",
  "最終 R3 與 LLM+RAG 基線相當；經 multi-seed 與 placebo 對照檢定，迭代於 SmartBugs 上無統計顯著之 F1 增益（詳見本節後段之迭代顯著性檢定）。"),
 # E — abstract: add rigor clause (already non-monotonic-honest)
 (64,
  "並經對抗式迭代後達 0.9128",
  "並經對抗式迭代後達 0.9128（R1=0.9164、R2=0.9060、R3=0.9128，呈非單調；經 multi-seed 與 placebo 對照檢定，迭代增益落於測量雜訊內）"),
 (67,
  "After adversarial iteration, the F1-score reaches 0.9128 (R1=0.9164, R2=0.9060, R3=0.9128; non-monotonic).",
  "After adversarial iteration, the F1-score reaches 0.9128 (R1=0.9164, R2=0.9060, R3=0.9128; non-monotonic); a multi-seed, placebo-controlled test confirms this iteration gain falls within measurement noise."),
]

ok, fail = 0, []
for idx, old, new in EDITS:
    p = doc.paragraphs[idx]
    if old in p.text:
        set_text(p, p.text.replace(old, new))
        ok += 1
        print(f"[OK] P{idx} edited")
    else:
        fail.append(idx)
        print(f"[MISS] P{idx} — phrase not found")

# B — insert new significance-test subsection right after P468
NEW = ("（迭代顯著性檢定）為釐清上述輪間變化係真實效益抑或測量雜訊，本研究以 3 個隨機種子"
       "（42、7、123）對「處理組（注入 Blue Team 知識補丁）」與「placebo 對照組（執行相同迭代迴圈"
       "但不注入補丁）」各執行完整三輪迭代，且每次執行前均將知識庫重置為相同之乾淨基線以確保各次執行"
       "之獨立性。結果顯示：處理組最終 F1 為 0.9149±0.0109、placebo 對照組為 0.9201±0.0072，兩者之"
       "信賴帶明顯重疊，且 placebo 平均值甚至略高於處理組（圖 4-X 之誤差棒所示）。此外，同一種子"
       "（seed=42）下處理組與 placebo 之首輪 F1 理應完全相同（因起點知識庫與種子一致），實測卻相差"
       " 0.0102，直接量化了 LLM API 於固定 seed 下之非決定性雜訊地板。綜上，對抗式迭代之知識回饋於"
       " SmartBugs 上並未帶來統計顯著之 F1 增益，先前觀察到之輪間波動應歸因於測量雜訊而非真實學習"
       "效益。此一結論係以對照實驗嚴格確立，構成本研究對多代理迭代方法之誠實實證評估。")
anchor = doc.paragraphs[469]
newp = anchor.insert_paragraph_before(NEW)
try:
    newp.style = doc.paragraphs[468].style
except Exception:
    pass
print("[OK] inserted significance-test subsection before P469")

doc.save(SRC)
print(f"\nDONE: {ok} edits applied, misses={fail}; saved {SRC}")
