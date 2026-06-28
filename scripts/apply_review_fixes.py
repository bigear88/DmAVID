#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""依審查報告 20260628-v1 修正論文（資料以 /home/curtis/DmAVID 為準）。
多數參考文獻前幾輪已修；本腳本處理仍待修的 6 處 + seed_placebo std。"""
import shutil
from docx import Document
SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_review.docx"
shutil.copyfile(SRC, BAK); print("backup:", BAK)
doc = Document(SRC)
paras = doc.paragraphs


def set_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)


# (段落索引, 舊子字串, 新子字串)
SUBS = [
 # P469 seed_placebo：母體 std → 樣本 std（審查 #新發現；專案原始 F1 算出 ddof=1=0.0133/0.0088）
 (469, "處理組最終 F1 為 0.9149±0.0109、placebo 對照組為 0.9201±0.0072",
       "處理組最終 F1 為 0.9149±0.0133、placebo 對照組為 0.9201±0.0088（n=3，樣本標準差）"),
 # #2 Chainalysis 標題 → 實際報告名
 (577, "Crypto hacking stolen funds 2024.", "The 2025 Crypto Crime Report."),
 # #4 Ganguli 完整副標題
 (581, "Red teaming language models to reduce harms. arXiv",
       "Red teaming language models to reduce harms: Methods, scaling behaviors, and lessons learned. arXiv"),
 # #13 Lin 補完整標題 "for static analysis"
 (593, "false positive mitigation. arXiv",
       "false positive mitigation for static analysis. arXiv"),
 # #19 RAG-SmartVuln 作者更正（審查確認 4 位；採越南名末字為排序鍵，內文 Nhu et al. 仍成立）
 (599, "Nhu, N. D. Q., Van, T. H., Trung, D. M., & Duy, P. T. (2025).",
       "Nhu, N. D. Q., Quan, L. M., Van, T. H., & Doan, T. M. (2025)."),
 # #32 Wei 2025a IEEE TSE 補 early-access 標註
 (613, "IEEE Transactions on Software Engineering.",
       "IEEE Transactions on Software Engineering. Advance online publication."),
]

for idx, old, new in SUBS:
    p = paras[idx]; t = p.text
    if old in t:
        set_text(p, t.replace(old, new)); print(f"[OK] P{idx}")
    else:
        print(f"[MISS] P{idx} 未找到: {old[:40]}")

doc.save(SRC)
print("SAVED")
