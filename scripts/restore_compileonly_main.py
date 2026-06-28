#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""依使用者決定：canonical 主文改回 compile-only BAK(2026-06-26) 單調 0.9158，
並明確說明採用此版本之缺點。從 bak_before_realpoc.docx 還原 option-C 迭代段落，
P468 強化為四點缺點，保留審查修正(P469 std、參考文獻)。"""
from docx import Document
SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_realpoc.docx"
cur = Document(SRC)
bak = Document(BAK)


def set_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell(cell, txt):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


# 1) 從備份逐段還原（option-C compile-only 主文）
RESTORE = [64, 67, 323, 406, 456, 460, 461, 467, 504, 532, 548, 558, 457, 459]
for i in RESTORE:
    set_text(cur.paragraphs[i], bak.paragraphs[i].text)
    print(f"[restore] P{i}")

# 2) 表格還原 0.9158
b6 = bak.tables[6].rows[5].cells; c6 = cur.tables[6].rows[5].cells
for j in range(len(c6)):
    set_cell(c6[j], b6[j].text)
b7 = bak.tables[7].rows[5].cells; c7 = cur.tables[7].rows[5].cells
for j in range(len(c7)):
    set_cell(c7[j], b7[j].text)
print("[restore] 表4-1/4-2 迭代列 -> 0.9158")

# 3) P468 強化為四點缺點（明確說明採用 compile-only 版本之缺點）
P468 = (
 "此一結果之核心意涵在於：在 solc 編譯把關（變體編譯通過即回饋）之知識回饋閉環下，多代理對抗式迭代於 SmartBugs 上呈現單調且正向之預測效能改善（三輪 F1 0.9103→0.9153→0.9158），支持 FN 課程學習閉環之設計理念——Blue Team 自漏判案例合成漏洞模式並經閉環回饋注入 Student，三輪共 26 筆補丁推動 Recall 由 0.9231 升至 0.9510。本研究選用此 compile-only 把關版本作為正式迭代之 canonical 結果，係因其能完整呈現迭代機制之學習軌跡與 FN 課程學習設計理念之可行性；惟須誠實揭露採用此版本之四項缺點與限制，以利讀者正確解讀其效能邊界。"
 "第一（對抗知識未經漏洞利用驗證，最根本之缺點）：compile-only 把關僅要求 Red Team 變體通過 solc 編譯即回饋知識庫，並未以 forge test 執行攻擊重放確認該變體之漏洞「確實可被利用」。因此知識庫可能納入「語法正確但實際不可利用」之偽對抗樣本，使所學漏洞模式之品質存疑，0.9061→0.9158 之 F1 增益宜視為效能上界而非保證可達之真實效能。"
 "第二（單調增益對驗證嚴格度高度敏感）：同一套迭代迴圈，若將把關改為嚴格之 forge test PoC 攻擊重放（要求漏洞確可被利用方回饋知識），則每輪可通過之有效對抗知識由 11/8/7 驟降至約 1 筆，學習訊號被大幅稀釋，三輪 F1 退化為非單調之 0.9164→0.9060→0.9128（最佳輪為 R1，R2 一度回落至基線），最終僅與單通道基線（0.9121）相當。顯示「單調逐輪改善」並非穩健性質，而高度依賴知識回饋閘門之寬鬆程度。"
 "第三（增益對隨機種子敏感、未達統計顯著）：以 3 個隨機種子（42、7、123）對處理組（注入補丁）與 placebo 對照組（不注入補丁）之嚴格檢定顯示，單調增益於跨種子下並不穩定，placebo 平均（0.9201）甚至略高於處理組（0.9149），且 +0.0037 之增益落於雜訊地板內、未達統計顯著（詳見後段之迭代顯著性檢定）。"
 "第四（可重現性侷限）：OpenAI API 於固定 seed 下仍具非決定性（同 seed=42 之首輪 F1 相差 0.0102），使單一 seed 之單調結果難以被他人完全重現。"
 "綜合而言，迭代之過程層面貢獻明確：FN 課程學習成立、知識庫動態擴充（補丁累積寫入 vulnerability_knowledge.json 並同步向量化至 ChromaDB），以及可解釋性深化——推理深度由平均 68 詞提升至 73.2 詞，修復建議生成率達 62.86%，程式碼行號引用精確率達 100%（此點將於第七節進一步驗證）；而其全資料集 F1 之增益則屬小幅，並具上述四項缺點，宜保守解讀。"
)
set_text(cur.paragraphs[468], P468)
print("[enhance] P468 四點缺點")

cur.save(SRC)
print("SAVED")
# 驗證
d = Document(SRC)
print("T6:", " | ".join(c.text.strip() for c in d.tables[6].rows[5].cells))
print("T7:", " | ".join(c.text.strip() for c in d.tables[7].rows[5].cells))
print("P457:", d.paragraphs[457].text)
print("P469 std 保留:", "0.9149±0.0133" in d.paragraphs[469].text)
print("P577 審查保留:", "North Korea drives" in d.paragraphs[577].text)
