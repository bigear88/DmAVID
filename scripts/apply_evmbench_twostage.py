#!/usr/bin/env python3
"""Drop the leaky 'Enhanced' (hint/gold-injected) config from the EVMbench section.
Rewrite as a clean two-stage story: detect-only 7.69% -> RAG+smart-preprocess 64.10%.
Also fix smart-preprocess technique description and make prose match the existing
time-stratified tables (which already use 64.10% as the 2024 boundary)."""
import shutil
from docx import Document

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_twostage.docx"
shutil.copyfile(SRC, BAK)
print("backup:", BAK)
doc = Document(SRC)

def set_text(p, txt):
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(txt)

# ---- 1. Delete the Enhanced row from 表 4-12 (Table index 17) ----
t = doc.tables[17]
target = None
for r in t.rows:
    if "Enhanced" in r.cells[0].text:
        target = r
        break
if target is not None:
    target._tr.getparent().remove(target._tr)
    print("[OK] removed Enhanced row from 表 4-12 (Table 17)")
else:
    print("[MISS] Enhanced row not found in Table 17")

# ---- 2. Full-paragraph rewrites ----
REWRITE = {
 477: ("Baseline 配置（純 LLM+RAG detect-only，不含任何 DeFi 專屬增強）之偵測率為 3/39=7.69%，"
       "僅能識別極少數漏洞。此一低偵測率反映出 SmartBugs 知識庫中之漏洞模式以傳統 Solidity 漏洞為主，"
       "對 DeFi 協議特有之邏輯缺陷覆蓋不足。在導入智能預處理機制（介面定義擷取、安全關鍵模組深度分析"
       "與匯入／繼承關係摘要）後，Smart preprocess 配置之偵測率大幅提升至 25/39=64.10%，且 Token 用量"
       "反由 124,812 降至 62,858（詳見第九款），顯示偵測率之提升主要源自輸入結構之最佳化而非推論成本之增加。"),
 494: ("第一階段為 SmartBugs 訓練前（pre-cutoff）資料集，DmAVID 取得 F1=0.9121，此為管線之最佳表現。"
       "第二階段為 EVMbench 2024 邊界（boundary）資料集，於 smart preprocess 配置下偵測率為 64.10%"
       "（25/39，詳見表 4-16）。第三階段為 2025 年後之 post-cutoff 資料集，包含 8 份審計報告中之 17 個漏洞，"
       "偵測率為 10/17=58.82%。"),
 499: ("智能預處理機制透過自動化之合約解析與結構化處理，將 EVMbench 之偵測率從 detect-only baseline 之 "
       "7.69% 大幅提升至 64.10%（25/39），如表 4-12 所示。此一增幅顯示 DeFi 合約之複雜結構（如 proxy 模式、"
       "library 引用、原始碼長度超出模型脈絡上限）對管線之原始輸入處理構成挑戰。智能預處理並非對合約原始碼"
       "直接截斷，而是擷取全部介面定義（函式簽章、事件與修飾子）、對安全關鍵模組（存取控制、資金流、外部呼叫）"
       "進行深度分析、並產生匯入／繼承關係圖摘要，藉此在不超出脈絡上限之前提下保留最具安全相關性之程式碼。"
       "值得注意的是，智能預處理於提升偵測率之同時亦降低了 Token 用量（由 detect-only 之 124,812 降至 62,858），"
       "原因在於僅保留安全關鍵模組與介面摘要、移除冗餘之相依程式碼，使送入 LLM 之有效輸入更為精簡，"
       "因而在提高偵測率之餘並未增加推論成本。"),
 503: ("SmartBugs（pre-cutoff，合約多源自 2018-2020 年）之 F1=0.9121，EVMbench 2024（boundary，smart "
       "preprocess）之偵測率 64.10%，post-cutoff 2025+ 之偵測率 58.82%。SmartBugs 上之高 F1 與真實未見 "
       "DeFi 合約上約六成偵測率之間之顯著落差，確認了預訓練記憶之存在——GPT-4.1-mini 之預訓練語料庫極可能"
       "包含 SmartBugs 中之早期合約，使其在 SmartBugs 上受益於記憶效應（此與 CodeBERT 於 EVMbench OOD 上 "
       "F1 僅 0.05 之崩跌互為佐證）。在真正未見之 EVMbench 上，2024 邊界（64.10%）與 2025+ post-cutoff"
       "（58.82%）之偵測率僅相差 5.28 個百分點，呈現平緩之衰退而非進一步崩跌，顯示框架之能力在跨越訓練截止點"
       "後大致穩定。其中 post-cutoff 17 個漏洞屬於現有 RAG 知識庫可涵蓋之傳統類型（重入攻擊、存取控制、"
       "狀態變數操控）者約佔 59%（10/17）且全數被成功偵測；反觀 EVMbench 2024 中涉及跨合約狀態依賴、代理模式"
       "與複雜 DeFi 協議邏輯之漏洞（如 Taiko 5 個、Forte 5 個均完全未被偵測）超出當前知識庫之覆蓋範圍，"
       "為各審計間偵測率差異之主要成因。"),
 504: ("此一分析之核心結論為：DmAVID 在 SmartBugs 上之高 F1（0.9121）確實包含預訓練記憶之貢獻，此為本研究"
       "坦誠承認之限制。然而，DmAVID 之真正貢獻並非來自 LLM 之原始偵測能力（LLM Base F1 僅 0.7474），"
       "而在於三項框架層面之增值：（1）RAG 知識檢索對 FPR 之大幅抑制（從 0.95 降至 0.24），（2）多代理對抗式"
       "迭代之貢獻——預測效能呈小幅且具雜訊之變化（三輪 F1 0.9164→0.9060→0.9128，最終 Recall 0.9510）"
       "與過程品質強化（3 筆學習補丁同步向量化寫入 ChromaDB），以及（3）多代理協作產生之高可解釋性輸出"
       "（修復建議率 62.86%）。即便在 post-cutoff 資料上（8 項審計，2025-01 至 2026-01），管線仍能維持 "
       "58.82% 之偵測率（10/17），其中 liquid-ron、blackhole、panoptic、tempo-feeamm 四項審計達 100% 偵測，"
       "顯示框架之泛化能力雖不完美但仍具實用價值。據此，宜將 DmAVID 面對「絕對未見」之真實 DeFi 合約之務實"
       "偵測率界定於約 58 至 64%（EVMbench 2024 邊界 64.10%、2025+ post-cutoff 58.82%），此即本框架真實之"
       "能力邊界，亦明確界定後續工作（DeFi 專屬邏輯漏洞之知識擴充）之方向。"),
}
for idx, txt in REWRITE.items():
    set_text(doc.paragraphs[idx], txt)
    print(f"[OK] P{idx} rewritten")

# ---- 3. Phrase replacements (captions / TOC / count word) ----
PHRASE = [
 (473, "進行三種配置之實驗", "進行兩種配置之實驗"),
 (474, "EVMbench 三配置偵測率比較", "EVMbench 兩配置偵測率比較"),
 (476, "EVMbench 三配置 DeFi 漏洞偵測率比較", "EVMbench 兩配置 DeFi 漏洞偵測率比較"),
 (128, "EVMbench 三配置偵測率比較", "EVMbench 兩配置偵測率比較"),
 (155, "EVMbench 三配置 DeFi 漏洞偵測率比較", "EVMbench 兩配置 DeFi 漏洞偵測率比較"),
]
for idx, old, new in PHRASE:
    p = doc.paragraphs[idx]
    if old in p.text:
        set_text(p, p.text.replace(old, new))
        print(f"[OK] P{idx} phrase replaced")
    else:
        print(f"[MISS] P{idx} phrase not found: {old}")

doc.save(SRC)
print("\nDONE. saved", SRC)
