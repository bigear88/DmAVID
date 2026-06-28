#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Phase 2: de-mechanize — remove .py script filenames (and the MUTATION_STRATEGIES
constant) from the thesis body so it reads as an academic paper, not a dev log.
Keeps function/JSON identifiers the professor retained (decide_early_stop(),
reload_dynamic_kb(), vulnerability_knowledge.json)."""
from docx import Document
SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
doc = Document(SRC)

def set_text(p, txt):
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)

REPS = [
 ("採用三種核心變體生成策略（12_red_team_generate.py MUTATION_STRATEGIES）。",
  "採用三種核心之紅隊突變策略。"),
 ("Foundry 可重現性驗證（13b_foundry_poc.py）作為 Stage 4",
  "Foundry 可重現性驗證作為 Stage 4"),
 ("（solc 編譯＋forge test PoC）由 13b_foundry_poc.py 實作並整合於自主協調器（20_coordinator_autonomous.py）之對抗迭代迴圈",
  "（solc 編譯＋forge test PoC）整合於自主協調器之對抗迭代迴圈"),
]
done = [0] * len(REPS)
for p in doc.paragraphs:
    t = p.text; new = t
    for i, (old, rep) in enumerate(REPS):
        if old in new:
            new = new.replace(old, rep); done[i] += 1
    if new != t:
        set_text(p, new)
for i, (old, _) in enumerate(REPS):
    print(("[OK]" if done[i] else "[MISS]"), old[:40])
doc.save(SRC)
# verify no .py left in body
import re
left = [(i, re.findall(r"[0-9A-Za-z_]+\.py", p.text)) for i, p in enumerate(Document(SRC).paragraphs) if re.search(r"[0-9A-Za-z_]+\.py", p.text)]
print("remaining .py paragraphs:", left if left else "NONE")
