#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Phase 3: align references to the professor-reviewed list.
- Add 3 missing entries (all cited in-text): Durieux 2020, Kevin & Yugopuspito 2025
  (SmartLLM, P248), SunWeb3Sec 2023 (DeFiHackLabs, Table 3-1), at correct A-Z positions.
- Fix ordering: Jiang before Jie (Jiang < Jie). (Wei 2025a/b already correct by title.)
"""
import copy, shutil
from docx import Document
from docx.oxml.ns import qn

SRC = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.docx"
BAK = "/mnt/d/OneDrive/DmAVID/DmAVID_20260622_v11.bak_before_refs.docx"
shutil.copyfile(SRC, BAK); print("backup:", BAK)
doc = Document(SRC)
paras = doc.paragraphs


def find(prefix):
    for p in paras:
        if p.text.strip().startswith(prefix):
            return p
    return None


def insert_ref_before(target, text):
    """Clone target <w:p> (keeps style/indent), reduce to one run, set text, insert before."""
    new_p = copy.deepcopy(target._p)
    # remove all runs
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # also remove hyperlinks if any
    for h in new_p.findall(qn("w:hyperlink")):
        new_p.remove(h)
    # build one run, copy rPr from target's first run if present
    src_r = target._p.find(qn("w:r"))
    new_r = copy.deepcopy(src_r) if src_r is not None else None
    if new_r is not None:
        for t in new_r.findall(qn("w:t")):
            new_r.remove(t)
        # clear any other children except rPr
        for ch in list(new_r):
            if ch.tag != qn("w:rPr"):
                new_r.remove(ch)
        wt = new_p.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        wt.text = text
        new_r.append(wt)
        new_p.append(new_r)
    else:
        from docx.text.paragraph import Paragraph
        Paragraph(new_p, target._parent).add_run(text)
    target._p.addprevious(new_p)


DURIEUX = ("Durieux, T., Ferreira, J. F., Abreu, R., & Cruz, R. (2020). Empirical review of "
           "automated analysis tools on 47,587 Ethereum smart contracts. In Proceedings of the "
           "42nd International Conference on Software Engineering (ICSE 2020), 530-541.")
KEVIN = ("Kevin, J., & Yugopuspito, P. (2025). SmartLLM: Smart contract vulnerability detection "
         "using large language models. In 2025 International Conference on Computer Sciences, "
         "Engineering, and Technology Innovation (ICoCSETI). IEEE. "
         "https://doi.org/10.1109/ICoCSETI63724.2025.11019687")
SUNWEB = ("SunWeb3Sec. (2023). DeFiHackLabs: Decentralized finance hack repository. GitHub. "
          "Retrieved from https://github.com/SunWeb3Sec/DeFiHackLabs")

# 1) insertions (skip if already present)
def ensure(text, before_prefix, label):
    if find(text[:24]):
        print(f"[skip] {label} already present"); return
    tgt = find(before_prefix)
    if tgt is None:
        print(f"[MISS] anchor not found for {label}: {before_prefix[:20]}"); return
    insert_ref_before(tgt, text)
    print(f"[OK] inserted {label} before {before_prefix[:18]}")

ensure(DURIEUX, "Feist, J., Grieco", "Durieux 2020")
ensure(KEVIN, "Li, Z., Dutta", "Kevin & Yugopuspito 2025")
ensure(SUNWEB, "Szabo, N. (1996)", "SunWeb3Sec 2023")

# 2) ordering fix: move Jiang before Jie
jie = find("Jie, W., Qiu")
jiang = find("Jiang, B., Liu")
if jie is not None and jiang is not None:
    jiang._p.getparent().remove(jiang._p)
    jie._p.addprevious(jiang._p)
    print("[OK] moved Jiang before Jie")
else:
    print("[MISS] Jie/Jiang anchors")

doc.save(SRC)

# verify order + presence
d2 = Document(SRC)
ps = d2.paragraphs
start = next(i for i, p in enumerate(ps) if p.text.strip() == "英文文獻")
ents = [p.text.strip() for p in ps[start+1:] if p.text.strip()]
print("\n英文條目數:", len(ents))
bad = [(a[:24], b[:24]) for a, b in zip(ents, ents[1:]) if a.lower() > b.lower()]
print("逆序對:", bad if bad else "無（已排序）")
for k in ["Durieux", "Kevin", "SunWeb3Sec"]:
    print(f"  {k} present:", any(e.startswith(k) for e in ents))
